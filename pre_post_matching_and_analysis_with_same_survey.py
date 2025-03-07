#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Created on Thu Dec 26 09:18:29 2024

@author: kieranmartin
"""
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches
from docx.shared import Pt
import json
import datetime
from fuzzywuzzy import fuzz 
from fuzzywuzzy import process
import re
from unidecode import unidecode
import os
os.chdir('/Users/kieranmartin/Documents/Qualtrics_API_Program')
from io import BytesIO
import QualAPI as qa
import requests
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re
from collections import Counter
# cd /Users/kieranmartin/Documents/Qualtrics_API_Program
# conda activate qualtrics_api_program
# spyder
# conda deactivate
pd.set_option('display.max_rows', None)  # Display all rows
pd.set_option('display.max_columns', None)  # Display all columns
pd.set_option('display.max_colwidth', None)  # Display full column content
pd.set_option('display.width', 1000)  # Set a large width for the display
###############################################################################
# Load Qualtrics credentials from a JSON file
with open('/Users/kieranmartin/Documents/Qualtrics_API_Program/qualtrics_credentials.txt') as f:
    creds = json.load(f)

# Extract client ID, secret, and data center from credentials
client_id = creds.get('ID')
client_secret = creds.get('Secret')
data_center = creds.get('DataCenter')
base_url = f'https://{data_center}.qualtrics.com'
###############################################################################
# if one survey and pre and post defined by cutoff date:
survey_name = "Lifestyle Survey -Fitness Programming (Adult)"

grant_type = 'client_credentials'
scope = 'read:surveys read:survey_responses'
data = qa.return_kwargs_as_dict(grant_type=grant_type, scope=scope)

# Get the bearer token
bearer_token_response = qa.get_token(base_url, client_id, client_secret, data)
token = bearer_token_response.get('access_token')

# Retrieve the list of all available surveys (to find our two IDs)
survey_list_df = qa.get_survey_list(base_url, token)
###############################################################################
# One survey with pre and post dataframes defined by cutoff date
###############################################################################
survey_id = qa.get_survey_id_by_name(survey_list_df, survey_name)

# Export survey responses and track the progress
response_export = qa.export_survey_responses(base_url, token, survey_id)
export_progress_id = response_export.get('result').get('progressId')

# Wait for export to complete and retrieve file ID for responses
file_id = qa.wait_for_export_completion(base_url, token, survey_id, export_progress_id)

# Download and process the survey responses
survey_responses = qa.get_survey_responses(base_url, token, survey_id, file_id)
responses_df = qa.extract_and_organize_responses(survey_responses)
responses_df = qa.filter_preview_responses(responses_df)

blocks_df = qa.get_block_data(base_url, survey_id, token)

# Retrieve survey questions and map them to response columns
survey_questions = qa.get_survey_questions(base_url, token, survey_id).get('result').get('questions')
question_df, question_values_df = qa.extract_column_data_types(
    survey_questions, 
    responses_df, 
    base_url, 
    token, 
    survey_id
)
question_df = qa.create_data_type_dictionary(question_df, question_values_df)

question_df = qa.reorder_question_df_with_normalized_ids(question_df, blocks_df)
#question_df = qa.reorder_question_df(question_df, blocks_df)
question_df = question_df.dropna()

question_values_df['question_value'] = question_values_df['question_value'].astype(str)
question_values_df['answer_id'] = question_values_df['answer_id'].astype(str)
question_values_df['question_id'] = question_values_df['question_id'].astype(str)

# Clean responses and retain only question columns
responses_df = qa.clean_responses(responses_df, question_df)
# Filter the DataFrame based on the conditions
filtered_df = responses_df[(responses_df['QID138'].notnull()) & (responses_df['progress'] > 82)]
# Include recordedDate in the filtered DataFrame
# Include recordedDate and ensure it is the first column in the DataFrame
columns_to_include = ['recordedDate'] + question_df['question_id'].tolist()
df = filtered_df.loc[:, columns_to_include].reset_index(drop=True)
df['recordedDate'] = pd.to_datetime(df['recordedDate'], format='ISO8601', errors='coerce')
df['recordedDate'] = df['recordedDate'].dt.strftime('%Y-%m-%d %H:%M:%S')


# Identify rows with all NaN values and filter them out
nan_mask = df.isna()
keep_mask = np.array(nan_mask.sum(axis=1) < len(df.columns))
df = df.loc[keep_mask].reset_index(drop=True)

##############################################################################
def single_clean_name(name_string: str) -> str:
    """
    Cleans a single name, returning a potentially empty or short string
    if it's 'test-like'. Does NOT decide whether to drop the row.
    """
    if not isinstance(name_string, str):
        return ""

    name = unidecode(name_string).lower()
    # Remove suffixes like jr, sr, etc.
    name = re.sub(r'(^|\s)(jr|sr|ii|iii|iv|v|vi)\b', '', name)

    # Remove non-alpha
    name = re.sub(r'[^a-z]', ' ', name)
    # Collapse multiple spaces
    name = re.sub(r'\s+', ' ', name).strip()

    # If it's in the known test set or extremely short,
    # it might not be a real name. We'll still return it here,
    # but we'll do the final logic in the row-based function.
    return name


def clean_names_rowwise(row, first_col='QID88_1', last_col='QID88_2'):
    """
    Cleans both first and last name in one row.
    Returns (cleaned_first, cleaned_last) OR (None, None)
    if BOTH are short/test-like OR match any of the specific (first, last) combos
    we want to drop.

    Then filter with df.dropna(subset=['QID88_1','QID88_2']) to drop those rows.
    """

    # A set of known "test-like" or short names
    known_test_set = {'fd', 'x', 'xx', 'xxx', 'test'}

    # A list of specific (first, last) combos that should cause the row to be dropped.
    # You can add as many as you need here.
    drop_combos = [
        ('melissa', 'otterbein testing'),
        ('melissa', 'otterbein'),
        ('jane', 'doe'),
        ('alicia', 'dixon ibarra')
    ]

    # 1) Clean each name individually
    raw_first = row.get(first_col, None)
    raw_last  = row.get(last_col, None)
    
    cleaned_first = single_clean_name(raw_first)
    cleaned_last  = single_clean_name(raw_last)

    # 2) Check if BOTH first and last are test-like or < 2 chars
    both_too_short_or_test = (
        (cleaned_first in known_test_set or len(cleaned_first) < 2) and
        (cleaned_last  in known_test_set or len(cleaned_last) < 2)
    )

    # 3) Check specific combos
    # e.g., if (cleaned_first, cleaned_last) == ('melissa', 'otterbein testing'), drop
    drop_specific_vals = (cleaned_first, cleaned_last) in drop_combos

    if both_too_short_or_test or drop_specific_vals:
        # Return (None, None) => we can later filter these out
        return None, None

    # Otherwise, keep the cleaned names
    return cleaned_first, cleaned_last


# Apply row-based cleaner.
df[['QID88_1', 'QID88_2']] = df.apply(clean_names_rowwise, axis=1, result_type='expand')
df = df.dropna(subset=['QID88_1','QID88_2'])

# Need to still handle duplicates:

pre_df = df[df['QID138'] == 1].copy()
post_df = df[df['QID138'] == 4].copy()
unsure_df = df[df['QID138'] == 5].copy()

def has_non_empty_qid104(val):
    """
    Returns True if QID104 is a non-empty list (or other data type).
    We treat [], None, np.nan, '' as 'empty'.
    Adjust this logic as needed if your column is stored differently.
    """
    # If val is None or NaN or empty string, it's empty
    if val is None or (isinstance(val, float) and pd.isna(val)) or val == '':
        return False
    # If val is a list and has length 0, it's empty
    if isinstance(val, list) and len(val) == 0:
        return False
    return True


def pick_valid_or_latest(subdf):
    """
    Given a sub-dataframe of duplicates (same QID88_1, QID88_2),
    pick the row that has a non-empty QID104 AND the latest recordedDate.
    If none have non-empty QID104, pick the row with the latest recordedDate anyway.
    """
    # 1) Convert recordedDate to datetime if not already
    if not pd.api.types.is_datetime64_any_dtype(subdf['recordedDate']):
        subdf['recordedDate'] = pd.to_datetime(subdf['recordedDate'])
    
    # 2) Filter to rows with non-empty QID104
    valid = subdf[subdf['QID104'].apply(has_non_empty_qid104)]
    
    if len(valid) > 0:
        # Among valid rows, pick the row with the latest date
        row_to_keep = valid.sort_values('recordedDate').iloc[-1:]
    else:
        # If no valid rows, pick the row with the latest date from the entire subdf
        row_to_keep = subdf.sort_values('recordedDate').iloc[-1:]
    
    return row_to_keep


def deduplicate_keep_valid_qid104(df):
    """
    Groups by (QID88_1, QID88_2) and keeps:
      - The latest recordedDate row among those that have a non-empty QID104
      - Otherwise, if all QID104 are empty, keep the row with the latest recordedDate.
    """
    # Group by first & last name
    grouped = df.groupby(['QID88_1','QID88_2'], group_keys=False)
    # Apply custom logic
    result = grouped.apply(pick_valid_or_latest)
    return result


pre_df = deduplicate_keep_valid_qid104(pre_df)
post_df = deduplicate_keep_valid_qid104(post_df)
unsure_df = deduplicate_keep_valid_qid104(unsure_df)

###############################################################################
# # columns to use if eventually want to incorperate more robust matching
# COLUMN_SPECS = {
#     # Name fields (strings)
#     'QID88_1': {'type': 'name', 'weight': 0.25},  # First Name
#     'QID88_2': {'type': 'name', 'weight': 0.25},  # Last Name
    
#     # Numeric field (allow ±1 for "close" match)
#     'QID3_TEXT': {'type': 'numeric', 'weight': 0.10},  # e.g. age
    
#     # Single-select fields (floats) => exact match
#     'QID44': {'type': 'single_select', 'weight': 0.10},  # living arrangement
#     'QID1716912176': {'type': 'single_select', 'weight': 0.05},  # region
#     'QID1716912177': {'type': 'single_select', 'weight': 0.05},  
#     'QID1716912178': {'type': 'single_select', 'weight': 0.05},
#     'QID1716912179': {'type': 'single_select', 'weight': 0.05},
#     'QID1716912180': {'type': 'single_select', 'weight': 0.05},
#     'QID1716912181': {'type': 'single_select', 'weight': 0.05},
#     'QID1716912182': {'type': 'single_select', 'weight': 0.05},
#     'QID1716912183': {'type': 'single_select', 'weight': 0.05},
    
#     # Multi-select fields => set-based similarity
#     # (each is something like ['3'], ['1','6'], or [] if not selected)
#     'QID55':  {'type': 'multi_select', 'weight': 0.10},  # race
#     'QID103': {'type': 'multi_select', 'weight': 0.05},  # gender
#     'QID104': {'type': 'multi_select', 'weight': 0.05},  # disability
# }

pre_df['orig_index_pre'] = pre_df.index
post_df['orig_index_post'] = post_df.index
unsure_df['orig_index_unsure'] = unsure_df.index

# 1) Create an empty "matches" DataFrame to store 100% matches from ANY pair.
matches_columns = [
    'pre_index', 'pre_first', 'pre_last',
    'post_index', 'post_first', 'post_last',
    'unsure_index', 'unsure_first', 'unsure_last',
    'match_score'
]
matches_df = pd.DataFrame(columns=matches_columns)


def exact_match_pre_post(pre_df, post_df):
    """
    Finds exact matches between pre_df and post_df on [QID88_1, QID88_2].
    Returns:
      - matched_df: with columns [pre_index, pre_first, pre_last,
                                  post_index, post_first, post_last, match_score=100]
      - leftover_pre: rows in pre_df that did NOT match
      - leftover_post: rows in post_df that did NOT match
    """
    # Make copies for safety, reset index
    ldf = pre_df.copy().reset_index(drop=True)
    rdf = post_df.copy().reset_index(drop=True)

    merged = pd.merge(
        ldf, rdf,
        how='inner',
        on=['QID88_1', 'QID88_2'],
        suffixes=('_pre','_post')  # for collisions if any
    )

    # Which indices matched?
    left_matched_indices = merged['orig_index_pre']
    right_matched_indices = merged['orig_index_post']

    # Leftovers
    leftover_pre = ldf[~ldf['orig_index_pre'].isin(left_matched_indices)]
    leftover_post = rdf[~rdf['orig_index_post'].isin(right_matched_indices)]

    # Build matched_df
    matched_df = pd.DataFrame({
        'pre_index':  merged['orig_index_pre'],
        'pre_first':  merged['QID88_1'],
        'pre_last':   merged['QID88_2'],
        'post_index': merged['orig_index_post'],
        'post_first': merged['QID88_1'],
        'post_last':  merged['QID88_2'],
        'unsure_index': None,
        'unsure_first': None,
        'unsure_last': None,
        'match_score': 100
    })

    return matched_df, leftover_pre, leftover_post


def exact_match_pre_unsure(pre_df, unsure_df):
    """
    Finds exact matches between pre_df and unsure_df on [QID88_1, QID88_2].
    Returns matched_df with pre* & unsure* columns, plus leftover dfs.
    """
    ldf = pre_df.copy().reset_index(drop=True)
    rdf = unsure_df.copy().reset_index(drop=True)

    merged = pd.merge(
        ldf, rdf,
        how='inner',
        on=['QID88_1', 'QID88_2'],
        suffixes=('_pre','_unsure')
    )

    left_matched_indices = merged['orig_index_pre']
    right_matched_indices = merged['orig_index_unsure']

    leftover_pre = ldf[~ldf['orig_index_pre'].isin(left_matched_indices)]
    leftover_unsure = rdf[~rdf['orig_index_unsure'].isin(right_matched_indices)]

    matched_df = pd.DataFrame({
        'pre_index':     merged['orig_index_pre'],
        'pre_first':     merged['QID88_1'],
        'pre_last':      merged['QID88_2'],
        'post_index':    None,
        'post_first':    None,
        'post_last':     None,
        'unsure_index':  merged['orig_index_unsure'],
        'unsure_first':  merged['QID88_1'],
        'unsure_last':   merged['QID88_2'],
        'match_score':   100
    })

    return matched_df, leftover_pre, leftover_unsure


def exact_match_post_unsure(post_df, unsure_df):
    """
    Finds exact matches between post_df and unsure_df on [QID88_1, QID88_2].
    Returns matched_df with post* & unsure* columns, plus leftover dfs.
    """
    ldf = post_df.copy().reset_index(drop=True)
    rdf = unsure_df.copy().reset_index(drop=True)

    merged = pd.merge(
        ldf, rdf,
        how='inner',
        on=['QID88_1','QID88_2'],
        suffixes=('_post','_unsure')
    )

    left_matched_indices = merged['orig_index_post']
    right_matched_indices = merged['orig_index_unsure']

    leftover_post = ldf[~ldf['orig_index_post'].isin(left_matched_indices)]
    leftover_unsure = rdf[~rdf['orig_index_unsure'].isin(right_matched_indices)]

    matched_df = pd.DataFrame({
        'pre_index':     None,
        'pre_first':     None,
        'pre_last':      None,
        'post_index':    merged['orig_index_post'],
        'post_first':    merged['QID88_1'],
        'post_last':     merged['QID88_2'],
        'unsure_index':  merged['orig_index_unsure'],
        'unsure_first':  merged['QID88_1'],
        'unsure_last':   merged['QID88_2'],
        'match_score':   100
    })

    return matched_df, leftover_post, leftover_unsure


###############################################################################
# 3) Perform EXACT matching in the order:
#    - pre vs. post
#    - leftover pre vs. unsure
#    - leftover post vs. unsure
#    Append each result to matches_df
###############################################################################
matched_pre_post, pre_leftover, post_leftover = exact_match_pre_post(pre_df, post_df)
matches_df = pd.concat([matches_df, matched_pre_post], ignore_index=True)

matched_pre_unsure, pre_leftover2, unsure_leftover1 = exact_match_pre_unsure(pre_leftover, unsure_df)
matches_df = pd.concat([matches_df, matched_pre_unsure], ignore_index=True)

matched_post_unsure, post_leftover2, unsure_leftover2 = exact_match_post_unsure(post_leftover, unsure_leftover1)
matches_df = pd.concat([matches_df, matched_post_unsure], ignore_index=True)

# Final leftover after all exact matches
pre_leftover_final = pre_leftover2
post_leftover_final = post_leftover2
unsure_leftover_final = unsure_leftover2

###############################################################################
# 4) Build a "leftover_partial_matches_df" with fuzzy partial matches
#    among the leftover pre/post/unsure sets.
#    We'll define a function to pick the best fuzzy match from B for each row in A.
###############################################################################
def name_match_score(firstA, lastA, firstB, lastB):
    """
    Weighted fuzzy match score (0..100).
    Adjust weights if last name should be more important.
    """
    score_first = fuzz.ratio(firstA, firstB)   # 0..100
    score_last = fuzz.ratio(lastA, lastB)      # 0..100
    # Weighted approach
    weight_first = 0.2
    weight_last = 0.8
    return (score_first * weight_first) + (score_last * weight_last)


def build_partial_matches(dfA, dfB, labelA, labelB):
    """
    For each row in dfA, find its best fuzzy match in dfB (if dfB not empty).
    Return a DataFrame with:
    [
      {labelA}_index, {labelA}_first, {labelA}_last,
      {labelB}_index, {labelB}_first, {labelB}_last,
      match_score
    ]
    """
    if dfB.empty:
        # If there's nothing in dfB, then all rows in A have no match
        return pd.DataFrame({
            f'{labelA}_index': dfA[f'orig_index_{labelA}'],
            f'{labelA}_first': dfA['QID88_1'],
            f'{labelA}_last':  dfA['QID88_2'],
            f'{labelB}_index': None,
            f'{labelB}_first': None,
            f'{labelB}_last':  None,
            'match_score': None
        })

    records = []
    b_records = dfB.to_dict(orient='records')  # faster than iterrows

    for _, rowA in dfA.iterrows():
        best_score = -1
        best_match = None
        for rowB in b_records:
            score = name_match_score(
                rowA['QID88_1'], rowA['QID88_2'],
                rowB['QID88_1'], rowB['QID88_2']
            )
            if score > best_score:
                best_score = score
                best_match = rowB

        rec = {
            f'{labelA}_index': rowA[f'orig_index_{labelA}'],
            f'{labelA}_first': rowA['QID88_1'],
            f'{labelA}_last':  rowA['QID88_2'],
            f'{labelB}_index': best_match[f'orig_index_{labelB}'] if best_match else None,
            f'{labelB}_first': best_match['QID88_1'] if best_match else None,
            f'{labelB}_last':  best_match['QID88_2'] if best_match else None,
            'match_score': best_score
        }
        records.append(rec)

    return pd.DataFrame(records)

# Build leftover partial matches for each pair:
partial_pre_post = build_partial_matches(pre_leftover_final, post_leftover_final, 'pre','post')
partial_pre_unsure = build_partial_matches(pre_leftover_final, unsure_leftover_final, 'pre','unsure')
partial_post_unsure = build_partial_matches(post_leftover_final, unsure_leftover_final, 'post','unsure')

# Combine them all if you want a single leftover partial dataframe:
leftover_partial_matches_df = pd.concat([partial_pre_post, partial_pre_unsure, partial_post_unsure],
                                        ignore_index=True)

###############################################################################
# Final Outputs:
###############################################################################
def confirm_partial_matches(
    row_indices_to_confirm,
    leftover_partial_matches_df,
    matches_df,
    pre_leftover_final,
    post_leftover_final,
    unsure_leftover_final
):
    """
    row_indices_to_confirm: list of integer indices (the index values in leftover_partial_matches_df)
    leftover_partial_matches_df: DataFrame of partial matches
    matches_df: your main matches DataFrame
    pre_leftover_final, post_leftover_final, unsure_leftover_final: leftover data
      that did not get an exact match.

    Returns updated dataframes:
      leftover_partial_matches_df, matches_df,
      pre_leftover_final, post_leftover_final, unsure_leftover_final
    with the newly confirmed matches removed from leftover and appended to matches_df.
    """

    # We will build a list of new rows to add to matches_df:
    new_matches = []

    # Sort the indices descending so that dropping them doesn't
    # interfere with subsequent indices (or we can just .loc them).
    # Typically safer to transform them into a list and process carefully.
    row_indices_to_confirm = sorted(row_indices_to_confirm, reverse=True)

    for idx in row_indices_to_confirm:
        if idx not in leftover_partial_matches_df.index:
            # If someone passes an invalid index, skip it or raise an error
            continue
        row = leftover_partial_matches_df.loc[idx]

        # Identify which side(s) are not None => that tells us if it's pre→post, pre→unsure, or post→unsure
        # We'll store them in a dict that aligns with matches_df columns.
        match_rec = {
            'pre_index':    None,
            'pre_first':    None,
            'pre_last':     None,
            'post_index':   None,
            'post_first':   None,
            'post_last':    None,
            'unsure_index': None,
            'unsure_first': None,
            'unsure_last':  None,
            'match_score':  row['match_score']
        }

        # If row['pre_index'] is not NaN => we have pre side
        if pd.notna(row.get('pre_index', None)):
            match_rec['pre_index'] = row['pre_index']
            match_rec['pre_first'] = row['pre_first']
            match_rec['pre_last']  = row['pre_last']

        # If row['post_index'] is not NaN => we have post side
        if pd.notna(row.get('post_index', None)):
            match_rec['post_index'] = row['post_index']
            match_rec['post_first'] = row['post_first']
            match_rec['post_last']  = row['post_last']

        # If row['unsure_index'] is not NaN => we have unsure side
        if pd.notna(row.get('unsure_index', None)):
            match_rec['unsure_index'] = row['unsure_index']
            match_rec['unsure_first'] = row['unsure_first']
            match_rec['unsure_last']  = row['unsure_last']

        # Append to new_matches
        new_matches.append(match_rec)

        # -------------------------------------------------------------
        # Remove matched rows from leftover final DataFrames
        #   - If we have a pre_index, remove that row from pre_leftover_final
        #   - If we have a post_index, remove that row from post_leftover_final
        #   - If we have an unsure_index, remove that row from unsure_leftover_final
        # -------------------------------------------------------------
        if pd.notna(match_rec['pre_index']):
            # Drop from pre_leftover_final (where orig_index_pre == match_rec['pre_index'])
            pre_leftover_final = pre_leftover_final[
                pre_leftover_final['orig_index_pre'] != match_rec['pre_index']
            ]

        if pd.notna(match_rec['post_index']):
            # Drop from post_leftover_final
            post_leftover_final = post_leftover_final[
                post_leftover_final['orig_index_post'] != match_rec['post_index']
            ]

        if pd.notna(match_rec['unsure_index']):
            # Drop from unsure_leftover_final
            unsure_leftover_final = unsure_leftover_final[
                unsure_leftover_final['orig_index_unsure'] != match_rec['unsure_index']
            ]

        # Also remove this row from leftover_partial_matches_df
        leftover_partial_matches_df = leftover_partial_matches_df.drop(idx)

    # -------------------------------------------------------------
    # Now append new_matches to matches_df
    # We can do this in one go if new_matches is not empty:
    # -------------------------------------------------------------
    if new_matches:
        new_matches_df = pd.DataFrame(new_matches)
        matches_df = pd.concat([matches_df, new_matches_df], ignore_index=True)

    return (
        leftover_partial_matches_df,
        matches_df,
        pre_leftover_final,
        post_leftover_final,
        unsure_leftover_final
    )


###############################################################################
# 2) Example usage
###############################################################################
# Examine leftover_partial_matches_df and determine which matches, less than 100%
# you want to include and hard code them here: This will unfortunatly have to be 
# re-checked everytime there is new data since I couldnt find a reliable threshold
# to cut off what is indefinitly a match and what is (like cases were a 90% match 
# score defitly wasnt a match but a 80% score was depending on how badly people
# mis-spelled there names)
leftover_partial_matches_df.loc[175]
leftover_partial_matches_df.loc[347]
leftover_partial_matches_df.loc[817]
leftover_partial_matches_df.loc[695]
leftover_partial_matches_df.loc[465]
# look good to you. Then do:

row_idxs = [175, 347, 817, 695, 465]
(
  leftover_partial_matches_df,
  matches_df,
  pre_leftover_final,
  post_leftover_final,
  unsure_leftover_final
) = confirm_partial_matches(
      row_idxs,
      leftover_partial_matches_df,
      matches_df,
      pre_leftover_final,
      post_leftover_final,
      unsure_leftover_final
)

# After this, those rows move into matches_df, and they're removed from leftover
# DataFrames & leftover_partial_matches_df. Next time the pipeline is run,
# they won't appear as leftovers again.
###############################################################################
def assign_match_ids(matches_df):
    """
    Resets the index of matches_df and assigns a unique match_id to each row.
    
    Parameters:
    matches_df (pd.DataFrame): DataFrame containing matching information.
    
    Returns:
    pd.DataFrame: Updated DataFrame with a new 'match_id' column.
    """
    matches_df = matches_df.reset_index(drop=True)
    matches_df['match_id'] = matches_df.index + 1
    return matches_df

matches_df = assign_match_ids(matches_df)


def generate_pre_matches(matches_df, pre_df, unsure_df):
    """
    Generates a DataFrame of pre-match rows by looking up each match in either pre_df or unsure_df.
    
    For each row in matches_df:
      - If 'pre_index' is not NA, fetch the matching row from pre_df where 'orig_index_pre' equals 'pre_index'.
      - Otherwise, fetch from unsure_df where 'orig_index_unsure' equals 'unsure_index'.
    
    Each matched row gets the corresponding 'match_id'. If no match is found,
    an empty row with the match_id is added.
    
    Parameters:
    matches_df (pd.DataFrame): DataFrame with matching information (must contain 'match_id', 'pre_index', and 'unsure_index').
    pre_df (pd.DataFrame): DataFrame containing the "pre" data, with column 'orig_index_pre'.
    unsure_df (pd.DataFrame): DataFrame containing the "unsure" data, with column 'orig_index_unsure'.
    
    Returns:
    pd.DataFrame: A concatenated DataFrame of pre-match rows with 'match_id' at the front.
    """
    pre_rows = []
    
    for _, match_row in matches_df.iterrows():
        if pd.notna(match_row['pre_index']):
            source_index = match_row['pre_index']
            subset_df = pre_df[pre_df['orig_index_pre'] == source_index]
        else:
            source_index = match_row['unsure_index']
            subset_df = unsure_df[unsure_df['orig_index_unsure'] == source_index]
        
        if len(subset_df) == 0:
            tmp = pd.Series(dtype='object')
            tmp['match_id'] = match_row['match_id']
            pre_rows.append(tmp)
        else:
            subset_df = subset_df.assign(match_id=match_row['match_id'])
            pre_rows.append(subset_df)
    
    pre_matches_df = pd.concat(pre_rows, ignore_index=True)
    cols = ['match_id'] + [c for c in pre_matches_df.columns if c != 'match_id']
    pre_matches_df = pre_matches_df[cols]
    return pre_matches_df

pre_matches_df = generate_pre_matches(matches_df, pre_df, unsure_df)


def generate_post_matches(matches_df, post_df, unsure_df):
    """
    Generates a DataFrame of post-match rows by looking up each match in either post_df or unsure_df.
    
    For each row in matches_df:
      - If 'post_index' is not NA, fetch the matching row from post_df where 'orig_index_post' equals 'post_index'.
      - Otherwise, fetch from unsure_df where 'orig_index_unsure' equals 'unsure_index'.
    
    Each matched row gets the corresponding 'match_id'. If no match is found,
    an empty row with the match_id is added.
    
    Parameters:
    matches_df (pd.DataFrame): DataFrame with matching information (must contain 'match_id', 'post_index', and 'unsure_index').
    post_df (pd.DataFrame): DataFrame containing the "post" data, with column 'orig_index_post'.
    unsure_df (pd.DataFrame): DataFrame containing the "unsure" data, with column 'orig_index_unsure'.
    
    Returns:
    pd.DataFrame: A concatenated DataFrame of post-match rows with 'match_id' at the front.
    """
    post_rows = []
    
    for _, match_row in matches_df.iterrows():
        if pd.notna(match_row['post_index']):
            source_index = match_row['post_index']
            subset_df = post_df[post_df['orig_index_post'] == source_index]
        else:
            source_index = match_row['unsure_index']
            subset_df = unsure_df[unsure_df['orig_index_unsure'] == source_index]
        
        if len(subset_df) == 0:
            tmp = pd.Series(dtype='object')
            tmp['match_id'] = match_row['match_id']
            post_rows.append(tmp)
        else:
            subset_df = subset_df.assign(match_id=match_row['match_id'])
            post_rows.append(subset_df)
    
    post_matches_df = pd.concat(post_rows, ignore_index=True)
    cols = ['match_id'] + [c for c in post_matches_df.columns if c != 'match_id']
    post_matches_df = post_matches_df[cols]
    return post_matches_df

post_matches_df = generate_post_matches(matches_df, post_df, unsure_df)


# Step 1: Assign Unique IDs
matches_df = matches_df.reset_index(drop=True)
matches_df['match_id'] = matches_df.index + 1

pre_rows = []

for i, match_row in matches_df.iterrows():
    # We'll store the single matched row from either pre_df or unsure_df
    if pd.notna(match_row['pre_index']):
        # Fetch from pre_df
        source_index = match_row['pre_index']
        subset_df = pre_df[pre_df['orig_index_pre'] == source_index]
    else:
        # Otherwise, fetch from unsure_df
        source_index = match_row['unsure_index']
        subset_df = unsure_df[unsure_df['orig_index_unsure'] == source_index]

    # We expect exactly 1 row, but just in case:
    if len(subset_df) == 0:
        # you could raise an error or create an empty row
        # for demonstration, let's just create an empty row
        # but typically you'd want to investigate why we didn't find a match
        tmp = pd.Series(dtype='object')
        tmp['match_id'] = match_row['match_id']
        pre_rows.append(tmp)
    else:
        # add match_id to each row in subset_df
        # (should be exactly 1, but let's be safe with .assign())
        subset_df = subset_df.assign(match_id=match_row['match_id'])
        pre_rows.append(subset_df)

pre_matches_df = pd.concat(pre_rows, ignore_index=True)

# Optionally reorder columns so match_id is at the front
cols = ['match_id'] + [c for c in pre_matches_df.columns if c != 'match_id']
pre_matches_df = pre_matches_df[cols]



post_rows = []

for i, match_row in matches_df.iterrows():
    if pd.notna(match_row['post_index']):
        # Real post row from post_df
        source_index = match_row['post_index']
        subset_df = post_df[post_df['orig_index_post'] == source_index]
    else:
        # Unsure is acting as "post"
        source_index = match_row['unsure_index']
        subset_df = unsure_df[unsure_df['orig_index_unsure'] == source_index]
    
    if len(subset_df) == 0:
        tmp = pd.Series(dtype='object')
        tmp['match_id'] = match_row['match_id']
        post_rows.append(tmp)
    else:
        subset_df = subset_df.assign(match_id=match_row['match_id'])
        post_rows.append(subset_df)

post_matches_df = pd.concat(post_rows, ignore_index=True)

# Optionally reorder columns
cols = ['match_id'] + [c for c in post_matches_df.columns if c != 'match_id']
post_matches_df = post_matches_df[cols]

########################### Report Generation #################################
regional_questions = [
    'Please select your SOAF Program: ',
    'Please select your SOAP Program: ',
    'Please select your SOEA Program:',
    'Please select your SOEE Program:',
    'Please select your SOLA Program:',
    'Please select your MENA Program:',
    'Please select your SONA Program:'
]

def process_response_column(responses_df, column):
    """
    Cleans and processes a multiple-choice or rank order response column by checking types 
    and handling lists, NaNs, and other cases as in the original structure.
    
    Parameters:
    responses_df (pd.DataFrame): Survey responses DataFrame.
    column (str): Column name for the multiple-choice or rank order question.
    
    Returns:
    list: Processed list of responses with 'NULL' for NaNs.
    """
    new_list = []
    
    # Check if the column is of object type
    if responses_df[column].dtype == 'object':
        for response in responses_df[column]:
            if isinstance(response, list):
                # Add list responses directly
                new_list += response
            elif isinstance(response, float) and np.isnan(response):
                new_list.append('NULL')
            elif isinstance(response, str) and response.isdigit():
                new_list.append(response)  # Keep numeric string values
            else:
                # Log unexpected values for detailed debugging
                print(f"Unexpected value in {column}: {response} (Type: {type(response)})")
                new_list.append('NULL')
    
    elif responses_df[column].dtype in ['float', 'int64', 'int']:  # Catch int64 as well
        for response in responses_df[column]:
            if isinstance(response, float) and np.isnan(response):
                new_list.append('NULL')
            else:
                new_list.append(str(int(response)))  # Convert to int and then to string
    else:
        print(f"Problems with conversion in column {column} (Unexpected data type: {responses_df[column].dtype})")
    
    return new_list


def process_multiple_choice_question(responses_df, col, question_values_df):
    """
    Processes a multiple-choice question to generate a frequency table
    with columns ['Code', 'Value', 'Count', 'Frequency'].
    """
    new_list = process_response_column(responses_df, col)
    freq_dist = Counter(new_list)
    freq_df = pd.DataFrame.from_dict(freq_dist, orient='index').reset_index()
    freq_df.columns = ['answer_id', 'N']
    
    current_values = question_values_df[question_values_df['question_id'] == col][['answer_id', 'question_value']]
    temp_df = pd.merge(current_values, freq_df, on='answer_id', how='outer').fillna({'N': 0, 'question_value': ''})
    temp_df['N'] = temp_df['N'].astype(int)
    temp_df['Pct'] = (temp_df['N'] / len(responses_df) * 100).round(1).astype(str) + '%'
    temp_df.columns = ['Code', 'Value', 'Count', 'Frequency']
    
    return temp_df


def process_numeric_question(responses_df, col, question_values_df):
    """
    Processes a numeric question to generate a frequency table
    with columns: ['Value', 'Count', 'Frequency'].
    
    Because question_values_df does not contain values for numeric questions,
    we skip merging.
    """
    new_list = process_response_column(responses_df, col)  # returns list of strings
    freq_dist = Counter(new_list)
    freq_df = pd.DataFrame.from_dict(freq_dist, orient='index').reset_index()
    freq_df.columns = ['Value', 'Count']

    # Calculate frequencies as percentages
    total_responses = len(responses_df)
    freq_df['Frequency'] = (freq_df['Count'] / total_responses * 100).round(1).astype(str) + '%'

    # Simply return freq_df, no merge required
    # with question_values_df since it's empty for numeric responses
    return freq_df[['Value', 'Count', 'Frequency']]


def combine_pre_post_tables_multiple_choice(pre_matches_df, post_matches_df, col, question_values_df):
    pre_result_df = process_multiple_choice_question(pre_matches_df, col, question_values_df)
    post_result_df = process_multiple_choice_question(post_matches_df, col, question_values_df)

    merged_df = pd.merge(
        pre_result_df,
        post_result_df,
        how='outer',
        on=['Code','Value'],
        suffixes=('_pre','_post')
    ).fillna({'Count_pre':0,'Count_post':0,'Frequency_pre':'0.0%','Frequency_post':'0.0%'})

    # Reorder columns
    merged_df = merged_df[['Code','Value','Count_pre','Frequency_pre','Count_post','Frequency_post']]
    merged_df.columns = ['Code','Value','Pre Count','Pre Frequency','Post Count','Post Frequency']
    
    # ---- CAST back to int so you don't get ".0" ----
    merged_df['Pre Count'] = merged_df['Pre Count'].astype(int)
    merged_df['Post Count'] = merged_df['Post Count'].astype(int)

    return merged_df


def combine_pre_post_tables_numeric(pre_matches_df, post_matches_df, col, question_values_df):
    pre_result_df = process_numeric_question(pre_matches_df, col, question_values_df)
    post_result_df = process_numeric_question(post_matches_df, col, question_values_df)
    
    merged_df = pd.merge(
        pre_result_df,
        post_result_df,
        how='outer',
        on='Value',
        suffixes=('_pre','_post')
    ).fillna({'Count_pre':0,'Count_post':0,'Frequency_pre':'0.0%','Frequency_post':'0.0%'})

    # Reorder columns
    merged_df = merged_df[['Value','Count_pre','Frequency_pre','Count_post','Frequency_post']]
    merged_df.columns = ['Value','Pre Count','Pre Frequency','Post Count','Post Frequency']
    
    # ---- CAST back to int ----
    merged_df['Pre Count'] = merged_df['Pre Count'].astype(int)
    merged_df['Post Count'] = merged_df['Post Count'].astype(int)
    
    return merged_df


def plot_pre_post_frequencies_mc(merged_df):
    """
    For a DataFrame with columns ['Code','Value','Pre Count','Pre Frequency',
                                 'Post Count','Post Frequency'],
    create an overlapping bar chart of Pre vs. Post frequencies (%).
    """
    # Convert percentage strings to floats
    pre_freq = merged_df['Pre Frequency'].str.rstrip('%').astype(float)
    post_freq = merged_df['Post Frequency'].str.rstrip('%').astype(float)
    
    x = np.arange(len(merged_df))
    width = 0.4
    
    fig, ax = plt.subplots(figsize=(7, 4.5))

    # Overlapping bars
    ax.bar(x - width/2, pre_freq, width=width, color='gray', alpha=0.7, label='Pre')
    ax.bar(x + width/2, post_freq, width=width, color='blue', alpha=0.7, label='Post')
    
    # Axis labels and ticks
    ax.set_xlabel('Code')       # or 'Value', depending on what you prefer
    ax.set_ylabel('Frequency (%)')
    ax.set_xticks(x)
    ax.set_xticklabels(merged_df['Code'])  # your code uses 'Code' for the x labels
    ax.set_ylim(0, 100)
    ax.legend()
    
    # Optionally annotate each bar with the numeric value (e.g., "52%")
    # The first half of ax.patches correspond to 'Pre' bars; the second half to 'Post'.
    for i, (freq_val, bar) in enumerate(zip(pre_freq, ax.patches[:len(x)])):
        ax.annotate(f'{freq_val}%', 
                    (bar.get_x() + bar.get_width()/2, bar.get_height() + 2),
                    ha='center', fontsize=9, color='gray')
    
    for i, (freq_val, bar) in enumerate(zip(post_freq, ax.patches[len(x):])):
        ax.annotate(f'{freq_val}%', 
                    (bar.get_x() + bar.get_width()/2, bar.get_height() + 2),
                    ha='center', fontsize=9, color='blue')
    
    # Grid lines for clarity
    ax.grid(axis='y', linestyle='--', alpha=0.7)
    
    plt.tight_layout()
    return fig


def plot_pre_post_frequencies_numeric(merged_df):
    """
    For a DataFrame with columns ['Value','Pre Count','Pre Frequency',
                                 'Post Count','Post Frequency'],
    create an overlapping bar chart of Pre vs. Post frequencies.
    """
    # Convert percentage strings like "42.1%" to float
    pre_freq = merged_df['Pre Frequency'].str.rstrip('%').astype(float)
    post_freq = merged_df['Post Frequency'].str.rstrip('%').astype(float)
    
    x = np.arange(len(merged_df))
    width = 0.4
    
    fig, ax = plt.subplots(figsize=(7, 4.5))

    ax.bar(x - width/2, pre_freq, width=width, color='gray', alpha=0.7, label='Pre')
    ax.bar(x + width/2, post_freq, width=width, color='blue', alpha=0.7, label='Post')
    
    ax.set_xlabel('Value')
    ax.set_ylabel('Frequency (%)')
    ax.set_xticks(x)
    ax.set_xticklabels(merged_df['Value'])
    ax.set_ylim(0, 100)
    ax.legend()
    
    # Annotate bars
    for i, (freq_val, bar) in enumerate(zip(pre_freq, ax.patches[:len(x)])):
        ax.annotate(f'{freq_val}%', 
                    (bar.get_x() + bar.get_width()/2, bar.get_height() + 2),
                    ha='center', fontsize=9, color='gray')
    for i, (freq_val, bar) in enumerate(zip(post_freq, ax.patches[len(x):])):
        ax.annotate(f'{freq_val}%', 
                    (bar.get_x() + bar.get_width()/2, bar.get_height() + 2),
                    ha='center', fontsize=9, color='blue')
    
    ax.grid(axis='y', linestyle='--', alpha=0.7)
    plt.tight_layout()
    return fig


def add_pre_post_table_to_doc(doc, merged_df):
    """
    Adds the pre/post frequency table to the Word doc.
    merged_df is either:
      - MC:   Code, Value, Pre Count, Pre Frequency, Post Count,  Post Frequency
      - Num:  Value, Pre Count, Pre Frequency, Post Count, Post Frequency
    """
    # Create table (header row + data rows)
    table = doc.add_table(rows=merged_df.shape[0] + 1, cols=merged_df.shape[1])
    table.style = 'Table Grid'
    
    # Header row
    for col_idx, col_name in enumerate(merged_df.columns):
        cell = table.cell(0, col_idx)
        cell.text = col_name
        cell.paragraphs[0].alignment = 1  # center
    
    # Data rows
    for row_idx in range(merged_df.shape[0]):
        for col_idx, val in enumerate(merged_df.iloc[row_idx]):
            cell = table.cell(row_idx + 1, col_idx)
            cell.text = str(val)
            # If the column is "Value," align left; else center
            if merged_df.columns[col_idx] == 'Value':
                cell.paragraphs[0].alignment = 0  # left
            else:
                cell.paragraphs[0].alignment = 1  # center

    return doc


def generate_pre_post_frequency_report(pre_matches_df, 
                                       post_matches_df, 
                                       question_df, 
                                       question_values_df, 
                                       regional_questions,
                                       survey_name="My Survey"):
    """
    Generates a Word doc comparing Pre vs. Post for both multiple-choice and numeric questions,
    in the order specified by question_df.
    """
    doc = Document()

    # Overview info (optional)
    doc.add_heading(f"{survey_name} - Pre vs. Post Report", level=1)
    doc.add_paragraph(f"Pre responses: N = {len(pre_matches_df)}", style='Normal')
    doc.add_paragraph(f"Post responses: N = {len(post_matches_df)}", style='Normal')
    doc.add_page_break()

    # Filter only the questions that are either MultipleChoice or Numeric (if needed)
    ordered_questions = question_df[question_df["data_type"].isin(['MultipleChoice', 'Numeric'])]

    # Iterate over question_df rows in the given order
    for idx, row in ordered_questions.iterrows():
        col = row['question_id']
        qname = row['question_name']
        qtext = row['question_text']
        mapped_text = f"{qname}: {qtext}"

        # Depending on the data type, call the corresponding processing functions
        if row['data_type'] == 'MultipleChoice':
            merged_df = combine_pre_post_tables_multiple_choice(pre_matches_df, post_matches_df, col, question_values_df)
            
            # (Optional) Filter out rows for regional questions
            if qtext in regional_questions:
                merged_df = merged_df[(merged_df['Pre Count'] != 0) | (merged_df['Post Count'] != 0)]
            
            # Add heading, table, and plot
            doc.add_heading(mapped_text, level=1)
            add_pre_post_table_to_doc(doc, merged_df)
            if len(merged_df) < 10:
                fig = plot_pre_post_frequencies_mc(merged_df)
                image_stream = BytesIO()
                plt.savefig(image_stream, format='png')
                plt.close(fig)
                image_stream.seek(0)
                doc.add_picture(image_stream, width=Inches(7), height=Inches(4.5))
        
        elif row['data_type'] == 'Numeric':
            merged_df = combine_pre_post_tables_numeric(pre_matches_df, post_matches_df, col, question_values_df)
            
            # (Optional) Filter out rows for regional questions
            if qtext in regional_questions:
                merged_df = merged_df[(merged_df['Pre Count'] != 0) | (merged_df['Post Count'] != 0)]
            
            # Add heading, table, and plot
            doc.add_heading(mapped_text, level=1)
            add_pre_post_table_to_doc(doc, merged_df)
            if len(merged_df) < 10:
                fig = plot_pre_post_frequencies_numeric(merged_df)
                image_stream = BytesIO()
                plt.savefig(image_stream, format='png')
                plt.close(fig)
                image_stream.seek(0)
                doc.add_picture(image_stream, width=Inches(7), height=Inches(4.5))
        
        # You can add additional "elif" blocks if you need to handle other data types

        doc.add_page_break()

    return doc


doc = generate_pre_post_frequency_report(
    pre_matches_df=pre_matches_df,
    post_matches_df=post_matches_df,
    question_df=question_df,
    question_values_df=question_values_df,
    regional_questions=regional_questions,
    survey_name="Lifestyle Survey"
)

doc.save("/Users/kieranmartin/Documents/Qualtrics_API_Program/fitness_report_test_bar_columns3.docx")









# regional_questions = [
#     'Please select your SOAF Program: ',
#     'Please select your SOAP Program: ',
#     'Please select your SOEA Program:',
#     'Please select your SOEE Program:',
#     'Please select your SOLA Program:',
#     'Please select your MENA Program:',
#     'Please select your SONA Program:'
# ]

# def process_response_column(responses_df, column):
#     """
#     Cleans and processes a multiple-choice or rank order response column by checking types 
#     and handling lists, NaNs, and other cases as in the original structure.
    
#     Parameters:
#     responses_df (pd.DataFrame): Survey responses DataFrame.
#     column (str): Column name for the multiple-choice or rank order question.
    
#     Returns:
#     list: Processed list of responses with 'NULL' for NaNs.
#     """
#     new_list = []
    
#     # Check if the column is of object type
#     if responses_df[column].dtype == 'object':
#         for response in responses_df[column]:
#             if isinstance(response, list):
#                 # Add list responses directly
#                 new_list += response
#             elif isinstance(response, float) and np.isnan(response):
#                 new_list.append('NULL')
#             elif isinstance(response, str) and response.isdigit():
#                 new_list.append(response)  # Keep numeric string values
#             else:
#                 # Log unexpected values for detailed debugging
#                 print(f"Unexpected value in {column}: {response} (Type: {type(response)})")
#                 new_list.append('NULL')
    
#     elif responses_df[column].dtype in ['float', 'int64', 'int']:  # Catch int64 as well
#         for response in responses_df[column]:
#             if isinstance(response, float) and np.isnan(response):
#                 new_list.append('NULL')
#             else:
#                 new_list.append(str(int(response)))  # Convert to int and then to string
#     else:
#         print(f"Problems with conversion in column {column} (Unexpected data type: {responses_df[column].dtype})")
    
#     return new_list


# def process_multiple_choice_question(responses_df, col, question_values_df):
#     """
#     Processes a multiple-choice question to generate a frequency table
#     with columns ['Code', 'Value', 'Count', 'Frequency'].
#     """
#     new_list = process_response_column(responses_df, col)
#     freq_dist = Counter(new_list)
#     freq_df = pd.DataFrame.from_dict(freq_dist, orient='index').reset_index()
#     freq_df.columns = ['answer_id', 'N']
    
#     current_values = question_values_df[question_values_df['question_id'] == col][['answer_id', 'question_value']]
#     temp_df = pd.merge(current_values, freq_df, on='answer_id', how='outer').fillna({'N': 0, 'question_value': ''})
#     temp_df['N'] = temp_df['N'].astype(int)
#     temp_df['Pct'] = (temp_df['N'] / len(responses_df) * 100).round(1).astype(str) + '%'
#     temp_df.columns = ['Code', 'Value', 'Count', 'Frequency']
    
#     return temp_df


# def process_numeric_question(responses_df, col, question_values_df):
#     """
#     Processes a numeric question to generate a frequency table
#     with columns: ['Value', 'Count', 'Frequency'].
#     """
#     new_list = process_response_column(responses_df, col)
#     freq_dist = Counter(new_list)
#     freq_df = pd.DataFrame.from_dict(freq_dist, orient='index').reset_index()
#     freq_df.columns = ['answer_id', 'N']
    
#     current_values = question_values_df[question_values_df['question_id'] == col][['answer_id', 'question_value']]
#     temp_df = pd.merge(current_values, freq_df, on='answer_id', how='outer').fillna({'N': 0, 'question_value': ''})
#     temp_df['N'] = temp_df['N'].astype(int)
#     temp_df['Pct'] = (temp_df['N'] / len(responses_df) * 100).round(1).astype(str) + '%'
    
#     # We rename to 'Value' (the numeric value/label), 'Count', 'Frequency'.
#     temp_df.columns = ['answer_id', 'Value', 'Count', 'Frequency']
    
#     # Drop 'answer_id' if you no longer need it:
#     temp_df = temp_df[['Value', 'Count', 'Frequency']]
    
#     return temp_df


# def combine_pre_post_tables_multiple_choice(pre_matches_df, post_matches_df, col, question_values_df):
#     pre_result_df = process_multiple_choice_question(pre_matches_df, col, question_values_df)
#     post_result_df = process_multiple_choice_question(post_matches_df, col, question_values_df)

#     merged_df = pd.merge(
#         pre_result_df,
#         post_result_df,
#         how='outer',
#         on=['Code','Value'],
#         suffixes=('_pre','_post')
#     ).fillna({'Count_pre':0,'Count_post':0,'Frequency_pre':'0.0%','Frequency_post':'0.0%'})

#     # Reorder columns
#     merged_df = merged_df[['Code','Value','Count_pre','Frequency_pre','Count_post','Frequency_post']]
#     merged_df.columns = ['Code','Value','Pre Count','Pre Frequency','Post Count','Post Frequency']
    
#     # ---- CAST back to int so you don't get ".0" ----
#     merged_df['Pre Count'] = merged_df['Pre Count'].astype(int)
#     merged_df['Post Count'] = merged_df['Post Count'].astype(int)

#     return merged_df


# def combine_pre_post_tables_numeric(pre_matches_df, post_matches_df, col, question_values_df):
#     pre_result_df = process_numeric_question(pre_matches_df, col, question_values_df)
#     post_result_df = process_numeric_question(post_matches_df, col, question_values_df)
    
#     merged_df = pd.merge(
#         pre_result_df,
#         post_result_df,
#         how='outer',
#         on='Value',
#         suffixes=('_pre','_post')
#     ).fillna({'Count_pre':0,'Count_post':0,'Frequency_pre':'0.0%','Frequency_post':'0.0%'})

#     # Reorder columns
#     merged_df = merged_df[['Value','Count_pre','Frequency_pre','Count_post','Frequency_post']]
#     merged_df.columns = ['Value','Pre Count','Pre Frequency','Post Count','Post Frequency']
    
#     # ---- CAST back to int ----
#     merged_df['Pre Count'] = merged_df['Pre Count'].astype(int)
#     merged_df['Post Count'] = merged_df['Post Count'].astype(int)
    
#     return merged_df


# def plot_pre_post_frequencies_mc_overlapped(merged_df):
#     """
#     Overlapping bar chart for multiple-choice question frequencies.
#     merged_df has columns: ['Code','Value','Pre Count','Pre Frequency','Post Count','Post Frequency'].
#     """
#     pre_freq = merged_df['Pre Frequency'].str.rstrip('%').astype(float)
#     post_freq = merged_df['Post Frequency'].str.rstrip('%').astype(float)

#     x = np.arange(len(merged_df))
#     width = 0.6  # adjust as you like

#     fig, ax = plt.subplots(figsize=(7, 4.5))

#     # 1) Pre bar (gray for all rows)
#     ax.bar(x, pre_freq, width=width, color='gray', alpha=0.5, label='Pre')

#     # 2) Post bar (blue except where Code == 'NULL', then red)
#     post_colors = ['red' if code == 'NULL' else 'blue' for code in merged_df['Code']]
#     ax.bar(x, post_freq, width=width, color=post_colors, alpha=0.8, label='Post')

#     ax.set_xlabel('Code')
#     ax.set_ylabel('Frequency (%)')
#     ax.set_xticks(x)
#     ax.set_xticklabels(merged_df['Code'])
#     ax.set_ylim(0, 100)
#     ax.legend()

#     # Annotate bars.
#     # The first len(x) patches are the 'Pre' bars; the second len(x) are the 'Post' bars.
#     for freq_val, bar in zip(pre_freq, ax.patches[:len(x)]):
#         ax.annotate(f'{freq_val}%',
#                     xy=(bar.get_x() + bar.get_width()/2, bar.get_height() + 2),
#                     ha='center', va='bottom', color='gray', fontsize=9)

#     for freq_val, bar in zip(post_freq, ax.patches[len(x):]):
#         # If the bar is red, make the text red to match (optional)
#         text_color = 'red' if bar.get_facecolor() == (1.0, 0.0, 0.0, 0.8) else 'blue'
#         ax.annotate(f'{freq_val}%',
#                     xy=(bar.get_x() + bar.get_width()/2, bar.get_height() + 2),
#                     ha='center', va='bottom', color=text_color, fontsize=9)

#     ax.grid(axis='y', linestyle='--', alpha=0.7)
#     plt.tight_layout()
#     return fig


# def plot_pre_post_frequencies_numeric_overlapped(merged_df):
#     """
#     Overlapping bar chart for numeric question frequencies.
#     merged_df has columns: ['Value','Pre Count','Pre Frequency','Post Count','Post Frequency'].
#     """
#     pre_freq = merged_df['Pre Frequency'].str.rstrip('%').astype(float)
#     post_freq = merged_df['Post Frequency'].str.rstrip('%').astype(float)

#     x = np.arange(len(merged_df))
#     width = 0.6

#     fig, ax = plt.subplots(figsize=(7, 4.5))

#     # Pre bars are always gray
#     ax.bar(x, pre_freq, width=width, color='gray', alpha=0.5, label='Pre')

#     # Post bars: red if Value == 'NULL', else blue
#     post_colors = ['red' if val == 'NULL' else 'blue' for val in merged_df['Value']]
#     ax.bar(x, post_freq, width=width, color=post_colors, alpha=0.8, label='Post')

#     ax.set_xlabel('Value')
#     ax.set_ylabel('Frequency (%)')
#     ax.set_xticks(x)
#     ax.set_xticklabels(merged_df['Value'])
#     ax.set_ylim(0, 100)
#     ax.legend()

#     # Annotate bars
#     for freq_val, bar in zip(pre_freq, ax.patches[:len(x)]):
#         ax.annotate(f'{freq_val}%',
#                     xy=(bar.get_x() + bar.get_width()/2, bar.get_height() + 2),
#                     ha='center', va='bottom', color='gray', fontsize=9)

#     for freq_val, bar in zip(post_freq, ax.patches[len(x):]):
#         text_color = 'red' if bar.get_facecolor() == (1.0, 0.0, 0.0, 0.8) else 'blue'
#         ax.annotate(f'{freq_val}%',
#                     xy=(bar.get_x() + bar.get_width()/2, bar.get_height() + 2),
#                     ha='center', va='bottom', color=text_color, fontsize=9)

#     ax.grid(axis='y', linestyle='--', alpha=0.7)
#     plt.tight_layout()
#     return fig



# def add_pre_post_table_to_doc(doc, merged_df):
#     """
#     Adds the pre/post frequency table to the Word doc.
#     merged_df is either:
#       - MC:   Code, Value, Pre Count, Pre Frequency, Post Count,  Post Frequency
#       - Num:  Value, Pre Count, Pre Frequency, Post Count, Post Frequency
#     """
#     # Create table (header row + data rows)
#     table = doc.add_table(rows=merged_df.shape[0] + 1, cols=merged_df.shape[1])
#     table.style = 'Table Grid'
    
#     # Header row
#     for col_idx, col_name in enumerate(merged_df.columns):
#         cell = table.cell(0, col_idx)
#         cell.text = col_name
#         cell.paragraphs[0].alignment = 1  # center
    
#     # Data rows
#     for row_idx in range(merged_df.shape[0]):
#         for col_idx, val in enumerate(merged_df.iloc[row_idx]):
#             cell = table.cell(row_idx + 1, col_idx)
#             cell.text = str(val)
#             # If the column is "Value," align left; else center
#             if merged_df.columns[col_idx] == 'Value':
#                 cell.paragraphs[0].alignment = 0  # left
#             else:
#                 cell.paragraphs[0].alignment = 1  # center

#     return doc


# def generate_pre_post_frequency_report(pre_matches_df, 
#                                        post_matches_df, 
#                                        question_df, 
#                                        question_values_df, 
#                                        regional_questions,
#                                        survey_name="My Survey"):
#     """
#     Generates a Word doc comparing Pre vs. Post for both multiple-choice and numeric questions.
#     """
#     # Identify question columns by data type
#     multiple_choice_columns = list(question_df.question_id[question_df["data_type"] == 'MultipleChoice'])
#     numeric_choice_columns = list(question_df.question_id[question_df["data_type"] == 'Numeric'])

#     doc = Document()

#     # Some overview info (optional)
#     doc.add_heading(f"{survey_name} - Pre vs. Post Report", level=1)
#     doc.add_paragraph(f"Pre responses: N = {len(pre_matches_df)}", style='Normal')
#     doc.add_paragraph(f"Post responses: N = {len(post_matches_df)}", style='Normal')
#     doc.add_page_break()

#     # Multiple-choice
#     for col in multiple_choice_columns:
#         qname = question_df.loc[question_df['question_id'] == col, 'question_name'].values[0]
#         qtext = question_df.loc[question_df['question_id'] == col, 'question_text'].values[0]
#         mapped_text = f"{qname}: {qtext}"

#         # Combine
#         merged_df = combine_pre_post_tables_multiple_choice(pre_matches_df, post_matches_df, col, question_values_df)
        
#         # (Optional) Filter out rows where both Pre Count & Post Count == 0 and question is regional
#         if qtext in regional_questions:
#             merged_df = merged_df[(merged_df['Pre Count'] != 0) | (merged_df['Post Count'] != 0)]
        
#         # Add heading
#         doc.add_heading(mapped_text, level=1)

#         # Insert table
#         add_pre_post_table_to_doc(doc, merged_df)

#         # Insert overlapping bar chart (if not too many categories)
#         if len(merged_df) < 10:
#             fig = plot_pre_post_frequencies_mc_overlapped(merged_df)
            
#             image_stream = BytesIO()
#             plt.savefig(image_stream, format='png')
#             plt.close(fig)
#             image_stream.seek(0)
            
#             doc.add_picture(image_stream, width=Inches(7), height=Inches(4.5))
        
#         doc.add_page_break()

#     # Numeric
#     for col in numeric_choice_columns:
#         qname = question_df.loc[question_df['question_id'] == col, 'question_name'].values[0]
#         qtext = question_df.loc[question_df['question_id'] == col, 'question_text'].values[0]
#         mapped_text = f"{qname}: {qtext}"

#         # Combine
#         merged_df = combine_pre_post_tables_numeric(pre_matches_df, post_matches_df, col, question_values_df)
        
#         # Optional filtering for regional
#         if qtext in regional_questions:
#             merged_df = merged_df[(merged_df['Pre Count'] != 0) | (merged_df['Post Count'] != 0)]
        
#         # Add heading
#         doc.add_heading(mapped_text, level=1)

#         # Insert table
#         add_pre_post_table_to_doc(doc, merged_df)

#         # Insert overlapping bar chart (if not too many categories)
#         if len(merged_df) < 10:
#             fig = plot_pre_post_frequencies_mc_overlapped(merged_df)
            
#             image_stream = BytesIO()
#             plt.savefig(image_stream, format='png')
#             plt.close(fig)
#             image_stream.seek(0)
            
#             doc.add_picture(image_stream, width=Inches(7), height=Inches(4.5))
        
#         doc.add_page_break()
    
#     return doc
# def generate_pre_post_frequency_report(pre_matches_df, 
#                                        post_matches_df, 
#                                        question_df, 
#                                        question_values_df, 
#                                        regional_questions,
#                                        survey_name="My Survey"):
#     """
#     Generates a Word doc comparing Pre vs. Post for both multiple-choice and numeric questions,
#     in the same block/question order that already exists in question_df.
#     """

#     # Instead of using sorted/filtered lists, we pull them directly from question_df
#     # (which is already in the desired block + question order).
#     multiple_choice_columns = question_df.loc[
#         question_df["data_type"] == "MultipleChoice", 
#         "question_id"
#     ].tolist()

#     numeric_choice_columns = question_df.loc[
#         question_df["data_type"] == "Numeric",
#         "question_id"
#     ].tolist()

#     doc = Document()

#     # Overview info (unchanged)
#     doc.add_heading(f"{survey_name} - Pre vs. Post Report", level=1)
#     doc.add_paragraph(f"Pre responses: N = {len(pre_matches_df)}", style='Normal')
#     doc.add_paragraph(f"Post responses: N = {len(post_matches_df)}", style='Normal')
#     doc.add_page_break()

#     # --------------------
#     # MULTIPLE-CHOICE LOOP
#     # --------------------
#     for col in multiple_choice_columns:
#         qname = question_df.loc[question_df['question_id'] == col, 'question_name'].values[0]
#         qtext = question_df.loc[question_df['question_id'] == col, 'question_text'].values[0]
#         mapped_text = f"{qname}: {qtext}"

#         # Combine
#         merged_df = combine_pre_post_tables_multiple_choice(pre_matches_df, post_matches_df, col, question_values_df)
        
#         # (Optional) Filter out rows where both Pre & Post Count == 0 and question is regional
#         if qtext in regional_questions:
#             merged_df = merged_df[(merged_df['Pre Count'] != 0) | (merged_df['Post Count'] != 0)]
        
#         # Add heading
#         doc.add_heading(mapped_text, level=1)

#         # Insert table
#         add_pre_post_table_to_doc(doc, merged_df)

#         # Insert overlapping bar chart (if not too many categories)
#         if len(merged_df) < 10:
#             fig = plot_pre_post_frequencies_mc_overlapped(merged_df)
            
#             image_stream = BytesIO()
#             plt.savefig(image_stream, format='png')
#             plt.close(fig)
#             image_stream.seek(0)
            
#             doc.add_picture(image_stream, width=Inches(7), height=Inches(4.5))
        
#         doc.add_page_break()

#     # -------------
#     # NUMERIC LOOP
#     # -------------
#     for col in numeric_choice_columns:
#         qname = question_df.loc[question_df['question_id'] == col, 'question_name'].values[0]
#         qtext = question_df.loc[question_df['question_id'] == col, 'question_text'].values[0]
#         mapped_text = f"{qname}: {qtext}"

#         # Combine
#         merged_df = combine_pre_post_tables_numeric(pre_matches_df, post_matches_df, col, question_values_df)
        
#         # Optional filtering for regional
#         if qtext in regional_questions:
#             merged_df = merged_df[(merged_df['Pre Count'] != 0) | (merged_df['Post Count'] != 0)]
        
#         # Add heading
#         doc.add_heading(mapped_text, level=1)

#         # Insert table
#         add_pre_post_table_to_doc(doc, merged_df)

#         # Insert overlapping bar chart (if not too many categories)
#         if len(merged_df) < 10:
#             fig = plot_pre_post_frequencies_mc_overlapped(merged_df)
            
#             image_stream = BytesIO()
#             plt.savefig(image_stream, format='png')
#             plt.close(fig)
#             image_stream.seek(0)
            
#             doc.add_picture(image_stream, width=Inches(7), height=Inches(4.5))
        
#         doc.add_page_break()
    
#     return doc



# def generate_pre_post_frequency_report(pre_matches_df, 
#                                        post_matches_df, 
#                                        question_df, 
#                                        question_values_df, 
#                                        regional_questions,
#                                        survey_name="My Survey"):
#     """
#     Generates a Word doc comparing Pre vs. Post for both multiple-choice and numeric questions,
#     following the same block + question order that already exists in `question_df`.
#     """

#     # 1) From question_df, gather question_ids (in the existing row order) 
#     #    for only "MultipleChoice" or "Numeric".
#     relevant_question_ids = question_df.loc[
#         question_df["data_type"].isin(["MultipleChoice", "Numeric"]),
#         "question_id"
#     ].tolist()

#     # 2) Create the doc (unchanged)
#     doc = Document()

#     # Overview info (unchanged)
#     doc.add_heading(f"{survey_name} - Pre vs. Post Report", level=1)
#     doc.add_paragraph(f"Pre responses: N = {len(pre_matches_df)}", style='Normal')
#     doc.add_paragraph(f"Post responses: N = {len(post_matches_df)}", style='Normal')
#     doc.add_page_break()

#     # 3) Loop over question IDs in the exact order they appear in question_df
#     for col in relevant_question_ids:
#         # Retrieve question name, text, and data_type from question_df
#         row_mask = question_df['question_id'] == col
#         qname     = question_df.loc[row_mask, 'question_name'].values[0]
#         qtext     = question_df.loc[row_mask, 'question_text'].values[0]
#         data_type = question_df.loc[row_mask, 'data_type'].values[0]

#         mapped_text = f"{qname}: {qtext}"

#         # 4) Combine pre/post tables depending on data_type
#         if data_type == "MultipleChoice":
#             merged_df = combine_pre_post_tables_multiple_choice(pre_matches_df, post_matches_df, col, question_values_df)
#         else:  # data_type == "Numeric"
#             merged_df = combine_pre_post_tables_numeric(pre_matches_df, post_matches_df, col, question_values_df)

#         # 5) (Optional) Filter out rows where both Pre & Post Count == 0 if question is "regional"
#         if qtext in regional_questions:
#             if 'Pre Count' in merged_df.columns and 'Post Count' in merged_df.columns:
#                 merged_df = merged_df[(merged_df['Pre Count'] != 0) | (merged_df['Post Count'] != 0)]

#         # 6) The rest of your code is unchanged: Add heading, table, chart, page break

#         # Add heading
#         doc.add_heading(mapped_text, level=1)

#         # Insert table
#         add_pre_post_table_to_doc(doc, merged_df)

#         # Insert overlapping bar chart (if not too many categories)
#         if len(merged_df) < 10:
#             # For multiple-choice
#             if data_type == "MultipleChoice":
#                 fig = plot_pre_post_frequencies_mc_overlapped(merged_df)
#             else:
#                 # For numeric
#                 fig = plot_pre_post_frequencies_mc_overlapped(merged_df)  
#                 # If you have a separate numeric overlapping function, you can call it here
#                 # e.g. plot_pre_post_frequencies_numeric_overlapped(merged_df)

#             image_stream = BytesIO()
#             plt.savefig(image_stream, format='png')
#             plt.close(fig)
#             image_stream.seek(0)
            
#             doc.add_picture(image_stream, width=Inches(7), height=Inches(4.5))

#         doc.add_page_break()
    
#     # 7) Return doc
#     return doc

# doc = generate_pre_post_frequency_report(
#     pre_matches_df=pre_matches_df,
#     post_matches_df=post_matches_df,
#     question_df=question_df,
#     question_values_df=question_values_df,
#     regional_questions=regional_questions,
#     survey_name="Lifestyle Survey"
# )

# doc.save("/Users/kieranmartin/Documents/Qualtrics_API_Program/fitness_report_test.docx")


# # Data for the chart
# categories = ['Bucket A', 'Bucket B', 'Bucket C', 'Bucket D', 'Bucket E']
# budgeted = [150, 130, 100, 90, 140]
# actual = [145, 125, 110, 80, 150]

# # Bar positions
# x = np.arange(len(categories))

# # Create the vertical bar chart
# fig, ax = plt.subplots(figsize=(7, 4.5))

# # Plot budgeted and actual values as overlapping vertical bars
# ax.bar(x - width/2, pre_freq, width=width, color='gray', alpha=0.7, label='Pre')
# ax.bar(x + width/2, post_freq, width=width, color='blue', label='Post')

# # Set axis labels, title, and ticks
# ax.set_ylabel('Frequency (%)')
# ax.set_xlabel('Code')  # Represents the categories (Buckets)
# ax.set_title('testing test test')
# ax.set_xticks(x)
# ax.set_xticklabels(merged_df['Value'])
# ax.set_ylim(0, 100)

# # Annotate each bar with its corresponding value
# for idx, (b_val, a_val) in enumerate(zip(budgeted, actual)):
#     ax.annotate(f'{b_val}', (idx, b_val + 2), ha='center', fontsize=9, color='gray')
#     ax.annotate(f'{a_val}', (idx, a_val + 2), ha='center', fontsize=9, color='blue')

# # Add legend
# ax.legend()

# # Add gridlines for clarity
# ax.grid(axis='y', linestyle='--', alpha=0.7)

# # Display the chart
# plt.tight_layout()
# plt.show()














