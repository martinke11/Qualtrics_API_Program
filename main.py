#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Created on Thu Aug 24 13:35:53 2023

@author: kieranmartin
"""
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches, Pt
import json
import datetime
import os
import requests
import qualtrics_API_functions as qa
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re
from collections import Counter

# Set the working directory
os.chdir('/Users/kieranmartin/Documents/SOI Data/Python/Qualtrics_API_Program')

# Inputs for wrapper function: Opening JSON file
with open('/Users/kieranmartin/Documents/SOI Data/Python/Qualtrics_API_Program/qualtrics_credentials.txt', 'r') as f:
    creds = json.load(f)

# Extract the clientId, clientSecret, and dataCenter
client_id = creds.get('ID')
client_secret = creds.get('Secret')
data_center = creds.get('DataCenter')

# Extract the Survey ID based on the name of the Survey
survey_name = 'Lifestyle Survey -Fitness Programming'

# Create the base URL
base_url = f'https://{data_center}.qualtrics.com'

# Set up the data for the token
grant_type = 'client_credentials'
scope = 'read:surveys read:survey_responses'
data = qa.def_data_t(grant_type=grant_type, scope=scope)

# Get the token
btkn = qa.get_token(base_url, client_id, client_secret, data)
# Extract the Bearer access token
tkn = btkn.get("access_token")

# Pull the list of surveys
survey_df = qa.get_surveys(base_url, tkn)

# Pull the indices associated with the Survey Name
inds = survey_df.loc[survey_df['name'] == survey_name].index[:]

if len(inds) > 1:
    print('Multiple Surveys Have the Same Name!!!')
elif len(inds) == 0:
    print('Cannot Find the Survey. Please Check that the Survey Name is Correct.')
elif len(inds) == 1:
    indx = inds[0]

# Extract the Survey Id
survey_id = survey_df.loc[indx, 'id']

# Get the Response Export
rx = qa.get_survey_rx(base_url, tkn, survey_id)
# Get the export progress ID
epid = rx.get('result').get('progressId')

# Wait until the status is 100% complete
fid = ""
while len(fid) == 0:
    # Check the progress
    rxp = qa.get_survey_rxp(base_url, tkn, survey_id, epid)

    # Check if the percent complete is 100
    if rxp.get('result').get('percentComplete') == 100:
        if rxp.get('result').get('status') == 'complete':
            # Get the file ID
            fid = rxp.get('result').get('fileId')
        else:
            print("Status: " + rxp.get('result').get('status'))

# Get the survey responses
survey = qa.get_survey(base_url, tkn, survey_id, fid)
# Convert to JSON
survey_json = survey.json()
# Get the list of responses
responses = survey_json.get('responses')

if len(responses) > 0:
    results = qa.organize_responses(responses)
else:
    print('There is no data')

# Do not keep survey preview responses
keep_mask = np.array(results['distributionChannel'] != 'preview')
results = results.loc[keep_mask, :]
# Reset the index
results = results.reset_index(drop=True)

# Filter responses with 100% progress and QID3 value of 2
#results = results[results['progress'] >= 100]
# results = results[results['QID3'] == 2]

# Get the survey questions
survey_qs = qa.get_survey_qs(base_url, tkn, survey_id)

# Map the Survey Qs so that they can match with column headers of Results
survey_qs_r = survey_qs.get('result')
q_dic = survey_qs_r.get('questions')

# Get the column data types
qdf, qvals = qa.extract_column_data_types(q_dic, results, base_url, tkn, survey_id)

# Get the different data types for reporting purposes
qdf = qa.dic_data_type(qdf, qvals)

# Using the QDF to extract only those columns with questions
df = results.loc[:, qdf['QID'].tolist()]

# Find and clean rows with NaN values
nan_mask = df.isna()
keep_mask = np.array(nan_mask.sum(axis=1) < len(df.columns))
results = results.loc[keep_mask, :]
df = df.loc[keep_mask, :]

# Reset the index
results = results.reset_index(drop=True)
df = df.reset_index(drop=True)

# Determine the percent of missing data by question
nan_mask = df.isna()
missing_data_n = nan_mask.sum(axis=0).tolist()
present_data_n = len(df) - np.array(missing_data_n)
missing_data_frc = np.array(missing_data_n) / len(df)
present_data_frc = 1 - missing_data_frc

# Create a DataFrame for missing data
tdic = {
    'QID': df.columns,
    'MissingDataN': missing_data_n,
    'PresentDataN': present_data_n,
    'MissingDataFrc': missing_data_frc,
    'PresentDataFrc': present_data_frc
}
missing_data_df = pd.DataFrame(tdic)

# Filter for FreeText columns
free_text_cols = qdf.QID[np.array(qdf.DataType == 'FreeText')]
mask = np.array(np.isin(missing_data_df['QID'], list(free_text_cols)))
text_miss_data_df = missing_data_df.loc[mask, :]

# Create frequency tables for multiple-choice questions
mc_cols = list(qdf.QID[np.array(qdf["DataType"] == 'MultipleChoice')])
result_dict = {}

for i in mc_cols:
    if isinstance(results[i], object):
        new_list = []
        for ii in results[i]:
            if isinstance(ii, list):
                new_list = new_list + ii
            elif isinstance(ii, float):
                if np.isnan(ii):
                    new_list = new_list + ['NULL'] 
                else:
                    new_list = new_list + [str(int(ii))]
            else:
                print('Problems with conversion')
    elif results[i].dtypes == 'float':
        new_list = []
        for ii in results[i]:
            if np.isnan(ii):
                new_list = new_list + ['NULL'] 
            else:
                new_list = new_list + [str(int(ii))]
    else:
        print('Problems with conversion')
        
         

    fdist = Counter(new_list)
    freq_df = pd.DataFrame.from_dict(fdist, orient='index').reset_index()
    freq_df.columns = ['Q_AnsId', 'N']

    # Merge with QVals
    cur_vals = qvals.loc[np.array(qvals['QID']) == i, ['Q_AnsId', 'Q_Value']]
    tdf = pd.merge(cur_vals, freq_df, on='Q_AnsId', how='outer')

    # Replace nans with zero
    Mask = np.array(np.isnan(tdf['N']))
    tdf.loc[Mask, 'N'] = 0
    tdf['N'] = tdf['N'].astype('int')
    # Replace nans with ''
    mask = np.array(tdf['Q_Value'].isna())
    tdf.loc[Mask, 'Q_Value'] = ''
    # Get the percent
    tdf['Pct'] = tdf['N'] / len(results) * 100
    tdf['Pct'] = tdf['Pct'].map('{0:.1f}'.format) + '%'
    
    # Set the names of the columns to however you want
    tdf.columns = ['Code', 'Value', 'Count', 'Frequency']
    # Code to calculate frequencies
    result_dict[i] = tdf
    

# Print the results from the dictionary
for column, result_df in result_dict.items():
    print(f"Column: {column}")
    print(result_df)
    print()

regional_questions = ['Please select your SOAF Program: ', 'Please select your SOAP Program: ',
'Please select your SOEA Program:', 'Please select your SOEE Program:',
'Please select your SOLA Program:', 'Please select your MENA Program:',
'Please select your SONA Program:']

# Save frequencies as a Word document
doc = Document()
section = doc.sections[0]
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1)
section.right_margin = Inches(1)

# Iterate through each column's results and add a table and bar chart
for original_column, result_df in result_dict.items():
    qname = qdf[qdf['QID'] == original_column]['QName'].values[0]
    qtext = qdf[qdf['QID'] == original_column]['QText'].values[0]
    regional_questions = [
        'Please select your SOAF Program:', 'Please select your SOAP Program:',
        'Please select your SOEA Program:', 'Please select your SOEE Program:',
        'Please select your SOLA Program:', 'Please select your MENA Program:',
        'Please select your SONA Program:'
    ]

    if qtext in regional_questions:
        filtered_result_df = result_df[result_df['Count'] != 0]
    else:
        filtered_result_df = result_df

    mapped_text = f"{qname}: {qtext}"
    doc.add_heading(mapped_text, level=1)

    # Add a table
    table = doc.add_table(rows=result_df.shape[0] + 1, cols=result_df.shape[1])
    table.style = 'Table Grid'

    column_alignment = [1] * result_df.shape[1]  # Center alignment
    value_column_index = result_df.columns.get_loc("Value")
    column_alignment[value_column_index] = 0  # Left align 'Value' column

    for row in table.rows:
        for cell in row.cells:
            cell.paragraphs[0].paragraph_format.alignment = 1
            cell.vertical_alignment = 1
            cell.width = Inches(2)

    for col_idx, column_name in enumerate(result_df.columns):
        cell = table.cell(0, col_idx)
        cell.text = column_name
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(14)
            paragraph.alignment = 1

    for row_idx in range(result_df.shape[0]):
        for col_idx, value in enumerate(result_df.iloc[row_idx]):
            cell = table.cell(row_idx + 1, col_idx)
            cell.text = str(value)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(12)
            cell.paragraphs[0].alignment = column_alignment[col_idx]

    # Create and save a bar chart
    fig, ax = plt.subplots(figsize=(7, 4.5))
    bar_width = 0.4

    frequency_values = filtered_result_df['Frequency'].astype(str).str.rstrip('%').astype(float)
    
    # Plotting the bar chart with the Frequency percentage
    usual_bar_color = 'dodgerblue'
    null_bar_color = 'crimson'
    
    bar_colors = [usual_bar_color if code != 'NULL' else null_bar_color for code in filtered_result_df['Code']]

    ax.bar(filtered_result_df['Code'], frequency_values, color=bar_colors, width=bar_width)
    ax.set_xlabel('Code')
    ax.set_ylabel('Frequency (%)')
    ax.set_xticks(filtered_result_df['Code'])
    ax.set_ylim(0, 100)

    for p, freq in zip(ax.patches, filtered_result_df['Frequency']):
        height = p.get_height()
        ax.annotate('{}'.format(freq),
                    xy=(p.get_x() + p.get_width() / 2, height),
                    xytext=(0, 3),  # 3 points vertical offset
                    textcoords="offset points",
                    ha='center', va='bottom', fontsize=10)



    # Save the bar chart as an image
    chart_filename = f'bar_chart_{original_column}.png'
    #plt.savefig(chart_filename)
    plt.close(fig)  # Close the figure
    doc.add_picture(chart_filename, width=Inches(7), height=Inches(4.5))
    doc.add_page_break()
    doc.add_page_break()

# Save the Word document
doc.save('/Users/kieranmartin/Documents/SOI Data/Python/Qualtrics_API_Program/Reports/lifestle.docx')




