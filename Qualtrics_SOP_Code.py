#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Created on Thu Aug 24 13:35:53 2023

@author: kieranmartin
"""
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches
from docx.shared import Pt
import json
import datetime
import os
os.chdir('/Users/kieranmartin/Documents/SOI Data/Python/Qualtrics_API_Program')
from io import BytesIO
import QualAPI as qa
import requests
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re
from collections import Counter
###############################################################################
# Load Qualtrics credentials from a JSON file
with open('/Users/kieranmartin/Documents/SOI Data/Python/Qualtrics_API_Program/qualtrics_credentials.txt') as f:
    creds = json.load(f)

# Extract client ID, secret, and data center from credentials
client_id = creds.get('ID')
client_secret = creds.get('Secret')
data_center = creds.get('DataCenter')
base_url = f'https://{data_center}.qualtrics.com'

# Define survey name and set up parameters for token request
survey_name = "Lifestyle Survey -Fitness Programming"
grant_type = 'client_credentials'
scope = 'read:surveys read:survey_responses'
data = qa.return_kwargs_as_dict(grant_type=grant_type, scope=scope)

# Get the bearer token
bearer_token_response = qa.get_token(base_url, client_id, client_secret, data)
token = bearer_token_response.get('access_token')

# Retrieve the list of surveys and find the survey ID
survey_list_df = qa.get_survey_list(base_url, token)
survey_id = qa.get_survey_id_by_name(survey_list_df, survey_name)

# Export survey responses and track the progress
response_export = qa.export_survey_responses(base_url, token, survey_id)
export_progress_id = response_export.get('result').get('progressId')

# Wait for export to complete and retrieve file ID for responses
file_id = qa.wait_for_export_completion(base_url, token, survey_id, export_progress_id)

# Download and process the survey responses
survey_responses = qa.get_survey_responses(base_url, token, survey_id, file_id)
responses_df = qa.extract_and_organize_responses(survey_responses)
responses_df = qa.filter_preview_responses(responses_df)

# Retrieve survey questions and map them to response columns
survey_questions = qa.get_survey_questions(base_url, token, survey_id).get('result').get('questions')
question_df, question_values_df = qa.extract_column_data_types(survey_questions, responses_df, base_url, token, survey_id)
question_df = qa.create_data_type_dictionary(question_df, question_values_df)

# Clean responses and retain only question columns
responses_df = qa.clean_responses(responses_df, question_df)
df = responses_df.loc[:, question_df['question_id'].tolist()]

# Identify rows with all NaN values to filter them out
nan_mask = df.isna()
keep_mask = np.array(nan_mask.sum(axis=1) < len(df.columns))
df = df.loc[keep_mask].reset_index(drop=True)

###############################################################################
# rest of code is for frequency analysis and report:
# use line below to filter responses_df by specific dates if needed:
# responses_df = qa.subset_by_date_range(responses_df, '2024-06-27', '2024-07-08')

def process_response_column(responses_df, column):
    """
    Cleans and processes a multiple-choice response column by checking types 
    and handling lists, NaNs, and other cases as in the original structure.
    
    Parameters:
    responses_df (pd.DataFrame): Survey responses DataFrame.
    column (str): Column name for the multiple-choice question.
    
    Returns:
    list: Processed list of responses with 'NULL' for NaNs.
    """
    new_list = []
    
    # Check if the column is of object type
    if isinstance(responses_df[column], object):
        for response in responses_df[column]:
            if isinstance(response, list):
                # Add the list of responses
                new_list += response
            elif isinstance(response, float):
                # Check if it's NaN and handle accordingly
                if np.isnan(response):
                    new_list.append('NULL')
                else:
                    new_list.append(str(int(response)))
            else:
                print('Problems with conversion')
                
    elif responses_df[column].dtype == 'float':
        for response in responses_df[column]:
            if np.isnan(response):
                new_list.append('NULL')
            else:
                new_list.append(str(int(response)))
    else:
        print('Problems with conversion')
    
    return new_list



regional_questions = ['Please select your SOAF Program: ', 'Please select your SOAP Program: ',
'Please select your SOEA Program:', 'Please select your SOEE Program:',
'Please select your SOLA Program:', 'Please select your MENA Program:',
'Please select your SONA Program:']

def generate_response_frequency(responses_df, question_df, question_values_df, regional_questions):
    """
    Generates frequency tables and plots for multiple-choice questions, 
    and saves them to a Word document.
    
    Parameters:
    responses_df (pd.DataFrame): Survey responses DataFrame.
    question_df (pd.DataFrame): DataFrame with question details, including 'question_id' and 'data_type'.
    question_values_df (pd.DataFrame): DataFrame with answer IDs and question values.
    regional_questions (list): List of questions that need regional filtering.

    Returns:
    Document: A Word document with tables and bar charts for each question.
    """
    result_dict = {}
    multiple_choice_columns = list(question_df.question_id[question_df["data_type"] == 'MultipleChoice'])

    for col in multiple_choice_columns:
        new_list = process_response_column(responses_df, col)
        # Create a frequency distribution of the new list
        freq_dist = Counter(new_list)
        freq_df = pd.DataFrame.from_dict(freq_dist, orient='index').reset_index()
        freq_df.columns = ['answer_id', 'N']  # Updated based on the column name in question_values_df
        
        # Create a new DataFrame with question_id, question_value, N, Pct
        current_values = question_values_df[question_values_df['question_id'] == col][['answer_id', 'question_value']]
        # Replace NaNs with empty strings for question values and with zero for counts
        temp_df = pd.merge(current_values, freq_df, on='answer_id', how='outer').fillna({'N': 0, 'question_value': ''})
        # Calculate the percentage
        temp_df['N'] = temp_df['N'].astype(int)
        temp_df['Pct'] = (temp_df['N'] / len(responses_df) * 100).round(1).astype(str) + '%'
        # Set the column names to desired labels
        temp_df.columns = ['Code', 'Value', 'Count', 'Frequency']
        # Store the resulting DataFrame in the dictionary
        result_dict[col] = temp_df
        
        # Print the frequency table for the current question
        # print(f"Frequency table for question {col}:\n{temp_df}\n")

    doc = Document()
    section = doc.sections[0]
    
    # Set the margins (all to 1 inch in this example)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    # Iterate through each column's results and add a table and a bar chart to the document
    for original_column, result_df in result_dict.items():
        # Find the corresponding question_name and question_text in question_df for the current question_id
        qname = question_df.loc[question_df['question_id'] == original_column, 'question_name'].values[0]
        qtext = question_df.loc[question_df['question_id'] == original_column, 'question_text'].values[0]
        # Combine question_name and question_text with a separator like ' - ' or ': '
        mapped_text = f"{qname}: {qtext}"
        
        # Add mapped_text as a heading
        doc.add_heading(mapped_text, level=1)
        filtered_result_df = result_df if qtext not in regional_questions else result_df[result_df['Count'] != 0]
        
        # Add a table
        table = doc.add_table(rows=filtered_result_df.shape[0] + 1, cols=filtered_result_df.shape[1])
        table.style = 'Table Grid'
        
        for col_idx, column_name in enumerate(filtered_result_df.columns):
            cell = table.cell(0, col_idx)
            cell.text = column_name
            cell.paragraphs[0].alignment = 1
        
        for row_idx in range(filtered_result_df.shape[0]):
            for col_idx, value in enumerate(filtered_result_df.iloc[row_idx]):
                cell = table.cell(row_idx + 1, col_idx)
                cell.text = str(value)
                cell.paragraphs[0].alignment = 1 if col_idx != 1 else 0  # Align 'Value' left, others center

        # The 10 here is how many question answers can be allowed in order for 
        # a bar plot to be created more than 10 wont fit in the visualization.
        # However all questions get tables generated. 
        if len(filtered_result_df) < 10:
            fig, ax = plt.subplots(figsize=(7, 4.5))
            # Extract frequency values for the bar chart
            frequency_values = filtered_result_df['Frequency'].str.rstrip('%').astype(float)
            # Plotting the bar chart with the Frequency percentage
            bar_colors = ['dodgerblue' if code != 'NULL' else 'crimson' for code in filtered_result_df['Code']]
            ax.bar(filtered_result_df['Code'], frequency_values, color=bar_colors, width=0.4)
            ax.set_xlabel('Code')
            ax.set_ylabel('Frequency (%)')
            ax.set_xticks(filtered_result_df['Code'])
            ax.set_ylim(0, 100)

            for p, freq in zip(ax.patches, filtered_result_df['Frequency']):
                ax.annotate(f'{freq}', (p.get_x() + p.get_width() / 2, p.get_height() + 3), ha='center', fontsize=10)
                
            # In order to not save the bar plots as indivuals picture files in folder 
            # use image stream:
            image_stream = BytesIO()
            plt.savefig(image_stream, format='png')
            plt.close(fig)
            image_stream.seek(0)
            doc.add_picture(image_stream, width=Inches(7), height=Inches(4.5))
            
            # If you WANT to save the indivual plots use code below instead:
            # Save the bar chart as an image
            # chart_filename = f'bar_chart_{original_column}.png'
            # plt.savefig(chart_filename)
            # plt.close(fig)  # Close the figure
            # doc.add_picture(chart_filename, width=Inches(7), height=Inches(4.5))

        doc.add_page_break()

    return doc


doc = generate_response_frequency(responses_df, question_df, question_values_df, regional_questions)
doc.save('/Users/kieranmartin/Documents/SOI Data/Python/Qualtrics_API_Program/Reports/lifestle3.docx')
