#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Created on Thu Aug 24 13:35:53 2023

@author: kieranmartin
"""
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches
from docx.shared import Pt
import json
import datetime
import os
os.chdir('/Users/kieranmartin/Documents/Qualtrics_API_Program')
from io import BytesIO
import QualAPI as qa
import requests
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re
from collections import Counter
###############################################################################
# Load Qualtrics credentials from a JSON file
with open('/Users/kieranmartin/Documents/Qualtrics_API_Program/qualtrics_credentials.txt') as f:
    creds = json.load(f)

# Extract client ID, secret, and data center from credentials
client_id = creds.get('ID')
client_secret = creds.get('Secret')
data_center = creds.get('DataCenter')
base_url = f'https://{data_center}.qualtrics.com'

# Define survey name and set up parameters for token request
survey_name = "Lifestyle Survey -Fitness Programming (Adult)"
#survey_name = "High 5 for Fitness Phase II- Teacher Survey"
grant_type = 'client_credentials'
scope = 'read:surveys read:survey_responses'
data = qa.return_kwargs_as_dict(grant_type=grant_type, scope=scope)

# Get the bearer token
bearer_token_response = qa.get_token(base_url, client_id, client_secret, data)
token = bearer_token_response.get('access_token')

# Retrieve the list of surveys and find the survey ID
survey_list_df = qa.get_survey_list(base_url, token)
survey_id = qa.get_survey_id_by_name(survey_list_df, survey_name)

# Export survey responses and track the progress
response_export = qa.export_survey_responses(base_url, token, survey_id)
export_progress_id = response_export.get('result').get('progressId')

# Wait for export to complete and retrieve file ID for responses
file_id = qa.wait_for_export_completion(base_url, token, survey_id, export_progress_id)

# Download and process the survey responses
survey_responses = qa.get_survey_responses(base_url, token, survey_id, file_id)
responses_df = qa.extract_and_organize_responses(survey_responses)
responses_df = qa.filter_preview_responses(responses_df)

# get survey question block info to order plots in order of questions in survey
# despite different data type handling
blocks_df = qa.get_block_data(base_url, survey_id, token)

# Retrieve survey questions and map them to response columns
survey_questions = qa.get_survey_questions(base_url, token, survey_id).get('result').get('questions')
question_df, question_values_df = qa.extract_column_data_types(survey_questions, responses_df, base_url, token, survey_id)
question_df = qa.create_data_type_dictionary(question_df, question_values_df)


question_df = qa.reorder_question_df_with_normalized_ids(question_df, blocks_df)
#question_df = qa.reorder_question_df(question_df, blocks_df)
question_df = question_df.dropna()


question_values_df['question_value'] = question_values_df['question_value'].astype(str)
question_values_df['answer_id'] = question_values_df['answer_id'].astype(str)
question_values_df['question_id'] = question_values_df['question_id'].astype(str)

# Clean responses and retain only question columns
responses_df = qa.clean_responses(responses_df, question_df)
df = responses_df.loc[:, question_df['question_id'].tolist()]

# Identify rows with all NaN values to filter them out
nan_mask = df.isna()
keep_mask = np.array(nan_mask.sum(axis=1) < len(df.columns))
df = df.loc[keep_mask].reset_index(drop=True)

# Filter question_df for FreeText columns and exclude those containing 'TEXT_'
free_text_columns = question_df[question_df.data_type == 'FreeText']

# Map question_id to question_text
question_id_to_text = free_text_columns.set_index('question_id')['question_text']

# Create text_df using question_id, then rename columns to question_text
text_df = df[question_id_to_text.index]
text_df.columns = text_df.columns.map(question_id_to_text)

###############################################################################
# rest of code is for frequency analysis and report:
# use line below to filter responses_df by specific dates if needed:
# responses_df = qa.subset_by_date_range(responses_df, '2024-06-27', '2024-07-08')
def process_response_column(responses_df, column):
    """
    Cleans and processes a multiple-choice or rank order response column by checking types 
    and handling lists, NaNs, and other cases as in the original structure.
    
    Parameters:
    responses_df (pd.DataFrame): Survey responses DataFrame.
    column (str): Column name for the multiple-choice or rank order question.
    
    Returns:
    list: Processed list of responses with 'NULL' for NaNs.
    """
    new_list = []
    
    # Check if the column is of object type
    if responses_df[column].dtype == 'object':
        for response in responses_df[column]:
            if isinstance(response, list):
                # Add list responses directly
                new_list += response
            elif isinstance(response, float) and np.isnan(response):
                new_list.append('NULL')
            elif isinstance(response, str) and response.isdigit():
                new_list.append(response)  # Keep numeric string values
            else:
                # Log unexpected values for detailed debugging
                print(f"Unexpected value in {column}: {response} (Type: {type(response)})")
                new_list.append('NULL')
    
    elif responses_df[column].dtype in ['float', 'int64', 'int']:  # Catch int64 as well
        for response in responses_df[column]:
            if isinstance(response, float) and np.isnan(response):
                new_list.append('NULL')
            else:
                new_list.append(str(int(response)))  # Convert to int and then to string
    else:
        print(f"Problems with conversion in column {column} (Unexpected data type: {responses_df[column].dtype})")
    
    return new_list


def process_multiple_choice_question(responses_df, col, question_values_df):
    """
    Processes a multiple-choice question to generate a frequency table.
    """
    new_list = process_response_column(responses_df, col)
    freq_dist = Counter(new_list)
    freq_df = pd.DataFrame.from_dict(freq_dist, orient='index').reset_index()
    freq_df.columns = ['answer_id', 'N']
    
    current_values = question_values_df[question_values_df['question_id'] == col][['answer_id', 'question_value']]
    temp_df = pd.merge(current_values, freq_df, on='answer_id', how='outer').fillna({'N': 0, 'question_value': ''})
    temp_df['N'] = temp_df['N'].astype(int)
    temp_df['Pct'] = (temp_df['N'] / len(responses_df) * 100).round(1).astype(str) + '%'
    temp_df.columns = ['Code', 'Value', 'Count', 'Frequency']
    
    return temp_df


def process_numeric_question(responses_df, col, question_values_df):
    """
    Processes a numeric question to generate a frequency table
    with columns: ['Value', 'Count', 'Frequency'].
    
    Because question_values_df does not contain values for numeric questions,
    we skip merging.
    """
    new_list = process_response_column(responses_df, col)  # returns list of strings
    freq_dist = Counter(new_list)
    freq_df = pd.DataFrame.from_dict(freq_dist, orient='index').reset_index()
    freq_df.columns = ['Value', 'Count']

    # Calculate frequencies as percentages
    total_responses = len(responses_df)
    freq_df['Frequency'] = (freq_df['Count'] / total_responses * 100).round(1).astype(str) + '%'

    # Simply return freq_df, no merge required
    # with question_values_df since it's empty for numeric responses
    return freq_df[['Value', 'Count', 'Frequency']]


regional_questions = [
    'Please select your SOAF Program: ',
    'Please select your SOAP Program: ',
    'Please select your SOEA Program:',
    'Please select your SOEE Program:',
    'Please select your SOLA Program:',
    'Please select your MENA Program:',
    'Please select your SONA Program:'
]

def generate_response_frequency(responses_df, question_df, question_values_df, regional_questions):
    """
    Generates frequency tables and plots for multiple-choice and numeric-choice questions with regional question filtering,
    and saves them to a Word document.
    
    Parameters:
    responses_df (pd.DataFrame): Survey responses DataFrame.
    question_df (pd.DataFrame): DataFrame with question details, including columns such as 'question_id', 'data_type',
                                'question_text', 'question_name', 'Block Name', and 'question_type'.
    question_values_df (pd.DataFrame): DataFrame with answer IDs and question values.
    regional_questions (list): List of question texts that need regional filtering.
    
    Returns:
    Document: A Word document with tables and bar charts for each applicable question.
    """
    from docx import Document
    from docx.shared import Inches
    import pandas as pd

    # Create a new Word document
    doc = Document()

    # Survey metadata
    N = len(responses_df)
    number_of_responses_text = f"N = {N} responses"
    total_questions = len(question_df)
    question_types = question_df['question_type'].unique()

    # Convert `endDate` to datetime and determine first/last response dates and survey duration
    responses_df['endDate'] = pd.to_datetime(responses_df['endDate'], format='%Y-%m-%dT%H:%M:%SZ')
    first_response_date = responses_df['endDate'].min().strftime('%Y-%m-%d')
    latest_response_date = responses_df['endDate'].max().strftime('%Y-%m-%d')
    survey_duration = (responses_df['endDate'].max() - responses_df['endDate'].min()).days

    # Add survey information to the document (survey_name should be defined elsewhere)
    doc.add_paragraph(survey_name, style='Normal')
    doc.add_paragraph(number_of_responses_text, style='Normal')
    doc.add_paragraph(
        f"Date of first survey response: {first_response_date}, "
        f"Date of most recently submitted survey response: {latest_response_date}",
        style='Normal'
    )
    doc.add_paragraph(f"Survey duration: {survey_duration} days", style='Normal')
    doc.add_paragraph(f"Number of questions: {total_questions}", style='Normal')
    doc.add_paragraph(f"Question types: {', '.join(question_types)}", style='Normal')
    doc.add_page_break()

    # Set custom margins
    section = doc.sections[0]
    section.top_margin = section.bottom_margin = section.left_margin = section.right_margin = Inches(1)

    # Iterate over question_df in its index order.
    # Process only questions with data types MultipleChoice or Numeric (ignoring FreeText questions).
    current_block = None
    for _, row in question_df.iterrows():
        if row['data_type'] not in ['MultipleChoice', 'Numeric']:
            continue  # Skip FreeText or other non-plottable question types

        # Group questions by block. When the block changes, add a new heading.
        block_name = row['Block Name']
        if block_name != current_block:
            doc.add_heading(block_name, level=1)
            current_block = block_name

        qid = row['question_id']
        qtext = row['question_text']
        qname = row['question_name']

        # Process the question based on its type.
        if row['data_type'] == 'MultipleChoice':
            result_df = process_multiple_choice_question(responses_df, qid, question_values_df)
        elif row['data_type'] == 'Numeric':
            result_df = process_numeric_question(responses_df, qid, question_values_df)

        # If this question requires regional filtering, remove zero-count responses.
        if qtext in regional_questions:
            result_df = result_df[result_df['Count'] != 0]

        # Add a heading for the question and insert its table/chart.
        mapped_text = f"{qname}: {qtext}"
        doc.add_heading(mapped_text, level=1)
        add_table_and_chart_to_doc(doc, result_df, is_numeric=(row['data_type'] == 'Numeric'))

    return doc



def add_table_and_chart_to_doc(doc, result_df, is_numeric=False):
    """
    Adds a frequency table and bar chart to the document for each question.
    For numeric questions, the chart and table use 'Value' instead of 'Code'.
    
    Parameters:
    doc (Document): The Word document being generated.
    result_df (pd.DataFrame): DataFrame containing the frequency table for the question.
    is_numeric (bool): Whether the question is numeric or multiple-choice.
    """
    # Determine column labels based on whether the question is numeric or multiple-choice
    code_column = 'Value' if is_numeric else 'Code'
    chart_label = 'Value' if is_numeric else 'Code'
    
    # Rename columns for the table if necessary
    display_df = result_df.rename(columns={'Code': code_column}) if is_numeric else result_df

    # Create and style table in the document
    table = doc.add_table(rows=display_df.shape[0] + 1, cols=display_df.shape[1])
    table.style = 'Table Grid'
    
    for col_idx, column_name in enumerate(display_df.columns):
        cell = table.cell(0, col_idx)
        cell.text = column_name
        cell.paragraphs[0].alignment = 1  # Center alignment
    
    for row_idx in range(display_df.shape[0]):
        for col_idx, value in enumerate(display_df.iloc[row_idx]):
            cell = table.cell(row_idx + 1, col_idx)
            cell.text = str(value)
            cell.paragraphs[0].alignment = 1 if col_idx != 1 else 0  # Align 'Value' left, others center

    # Create a bar chart if the result has fewer than 10 entries
    if len(display_df) < 10:
        fig, ax = plt.subplots(figsize=(7, 4.5))
        frequency_values = display_df['Frequency'].str.rstrip('%').astype(float)
        bar_colors = ['dodgerblue' if val != 'NULL' else 'crimson' for val in display_df[code_column]]
        ax.bar(display_df[code_column], frequency_values, color=bar_colors, width=0.4)
        ax.set_xlabel(chart_label)
        ax.set_ylabel('Frequency (%)')
        ax.set_xticks(display_df[code_column])
        ax.set_ylim(0, 100)

        for p, freq in zip(ax.patches, display_df['Frequency']):
            ax.annotate(f'{freq}', (p.get_x() + p.get_width() / 2, p.get_height() + 3), ha='center', fontsize=10)

        # Save chart as an image in memory and add to the document
        image_stream = BytesIO()
        plt.savefig(image_stream, format='png')
        plt.close(fig)
        image_stream.seek(0)
        doc.add_picture(image_stream, width=Inches(7), height=Inches(4.5))
    
    doc.add_page_break()
    
    return doc

# Add regional_questions parameter when calling the function
doc = generate_response_frequency(responses_df, question_df, question_values_df, regional_questions)
doc.save('/Users/kieranmartin/Documents/Qualtrics_API_Program/test_again20.docx')



