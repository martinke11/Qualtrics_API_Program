#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Created on Sun Dec 22 10:42:49 2024

@author: kieranmartin
"""
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches
from docx.shared import Pt
import json
import datetime
import os
os.chdir('/Users/kieranmartin/Documents/Qualtrics_API_Program')
from io import BytesIO
import QualAPI as qa
import requests
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re
from collections import Counter
###############################################################################
# Load Qualtrics credentials from a JSON file
with open('/Users/kieranmartin/Documents/Qualtrics_API_Program/qualtrics_credentials.txt') as f:
    creds = json.load(f)

# Extract client ID, secret, and data center from credentials
client_id = creds.get('ID')
client_secret = creds.get('Secret')
data_center = creds.get('DataCenter')
base_url = f'https://{data_center}.qualtrics.com'
###############################################################################
# if one survey and pre and post defined by cutoff date:
survey_name = ""
# Instead of one survey_name, let's define two:
pre_survey = ""
post_survey = ""

grant_type = 'client_credentials'
scope = 'read:surveys read:survey_responses'
data = qa.return_kwargs_as_dict(grant_type=grant_type, scope=scope)

# Get the bearer token
bearer_token_response = qa.get_token(base_url, client_id, client_secret, data)
token = bearer_token_response.get('access_token')

# Retrieve the list of all available surveys (to find our two IDs)
survey_list_df = qa.get_survey_list(base_url, token)
###############################################################################
# One survey with pre and post dataframes defined by cutoff date
###############################################################################
survey_id = qa.get_survey_id_by_name(survey_list_df, survey_name)

# Export survey responses and track the progress
response_export = qa.export_survey_responses(base_url, token, survey_id)
export_progress_id = response_export.get('result').get('progressId')

# Wait for export to complete and retrieve file ID for responses
file_id = qa.wait_for_export_completion(base_url, token, survey_id, export_progress_id)

# Download and process the survey responses
survey_responses = qa.get_survey_responses(base_url, token, survey_id, file_id)
responses_df = qa.extract_and_organize_responses(survey_responses)
responses_df = qa.filter_preview_responses(responses_df)

# Retrieve survey questions and map them to response columns
survey_questions = qa.get_survey_questions(base_url, token, survey_id).get('result').get('questions')
question_df, question_values_df = qa.extract_column_data_types(
    survey_questions, 
    responses_df, 
    base_url, 
    token, 
    survey_id
)
question_df = qa.create_data_type_dictionary(question_df, question_values_df)

# Clean responses and retain only question columns
responses_df = qa.clean_responses(responses_df, question_df)
df = responses_df.loc[:, question_df['question_id'].tolist()]

# Identify rows with all NaN values and filter them out
nan_mask = df.isna()
keep_mask = np.array(nan_mask.sum(axis=1) < len(df.columns))
df = df.loc[keep_mask].reset_index(drop=True)

cutoff_date = "2024-01-01"  # <-- adjust as needed

# Make sure "recordedDate" is a datetime column
if 'recordedDate' in responses_df.columns:
    responses_df['recordedDate'] = pd.to_datetime(responses_df['recordedDate'])
else:
    # If for some reason your recorded date column is not in responses_df
    # or has a different name, adjust accordingly
    raise ValueError("No 'recordedDate' column found in responses_df.")

# Create the pre and post slices. (We usually do this *before* filtering columns 
# down to question IDs, so we have the original date info available.)
pre_df = responses_df[responses_df['recordedDate'] <= cutoff_date].copy()
post_df = responses_df[responses_df['recordedDate'] > cutoff_date].copy()

# (Optional) Clean each subset separately using your question_df
# if you need only question columns in each split:
pre_df = pre_df.loc[:, question_df['question_id'].tolist()]
post_df = post_df.loc[:, question_df['question_id'].tolist()]

# Re-check for NaN rows in each subset
nan_mask_pre = pre_df.isna()
keep_mask_pre = np.array(nan_mask_pre.sum(axis=1) < len(pre_df.columns))
pre_df = pre_df.loc[keep_mask_pre].reset_index(drop=True)

nan_mask_post = post_df.isna()
keep_mask_post = np.array(nan_mask_post.sum(axis=1) < len(post_df.columns))
post_df = post_df.loc[keep_mask_post].reset_index(drop=True)

###############################################################################
# 1) PRE-SURVEY:Defined as one seperate survey
###############################################################################
pre_survey_id = qa.get_survey_id_by_name(survey_list_df, pre_survey)

# Export pre-survey responses and track progress
pre_response_export = qa.export_survey_responses(base_url, token, pre_survey_id)
pre_export_progress_id = pre_response_export.get('result').get('progressId')

# Wait for export to complete
pre_file_id = qa.wait_for_export_completion(base_url, token, pre_survey_id, pre_export_progress_id)

# Download and process the pre-survey responses
pre_survey_responses = qa.get_survey_responses(base_url, token, pre_survey_id, pre_file_id)
pre_responses_df = qa.extract_and_organize_responses(pre_survey_responses)
pre_responses_df = qa.filter_preview_responses(pre_responses_df)

# Retrieve pre-survey questions and map them to response columns
pre_survey_questions = qa.get_survey_questions(base_url, token, pre_survey_id).get('result').get('questions')
pre_question_df, pre_question_values_df = qa.extract_column_data_types(
    pre_survey_questions, 
    pre_responses_df,
    base_url,
    token,
    pre_survey_id
)
pre_question_df = qa.create_data_type_dictionary(pre_question_df, pre_question_values_df)

# Clean pre-survey responses and retain only question columns
pre_responses_df = qa.clean_responses(pre_responses_df, pre_question_df)
pre_df = pre_responses_df.loc[:, pre_question_df['question_id'].tolist()]

# Identify rows with all NaN values to filter them out
nan_mask_pre = pre_df.isna()
keep_mask_pre = np.array(nan_mask_pre.sum(axis=1) < len(pre_df.columns))
pre_df = pre_df.loc[keep_mask_pre].reset_index(drop=True)

###############################################################################
# 2) POST-SURVEY:Defined as one seperate survey
###############################################################################
post_survey_id = qa.get_survey_id_by_name(survey_list_df, post_survey)

# Export post-survey responses and track progress
post_response_export = qa.export_survey_responses(base_url, token, post_survey_id)
post_export_progress_id = post_response_export.get('result').get('progressId')

# Wait for export to complete
post_file_id = qa.wait_for_export_completion(base_url, token, post_survey_id, post_export_progress_id)

# Download and process the post-survey responses
post_survey_responses = qa.get_survey_responses(base_url, token, post_survey_id, post_file_id)
post_responses_df = qa.extract_and_organize_responses(post_survey_responses)
post_responses_df = qa.filter_preview_responses(post_responses_df)

# Retrieve post-survey questions and map them to response columns
post_survey_questions = qa.get_survey_questions(base_url, token, post_survey_id).get('result').get('questions')
post_question_df, post_question_values_df = qa.extract_column_data_types(
    post_survey_questions, 
    post_responses_df,
    base_url,
    token,
    post_survey_id
)
post_question_df = qa.create_data_type_dictionary(post_question_df, post_question_values_df)

# Clean post-survey responses and retain only question columns
post_responses_df = qa.clean_responses(post_responses_df, post_question_df)
post_df = post_responses_df.loc[:, post_question_df['question_id'].tolist()]

# Identify rows with all NaN values to filter them out
nan_mask_post = post_df.isna()
keep_mask_post = np.array(nan_mask_post.sum(axis=1) < len(post_df.columns))
post_df = post_df.loc[keep_mask_post].reset_index(drop=True)

###############################################################################
# Plotting and Report Generating: Uses pre_df and post_df created by either
# method above (1 survey or 2 seperate surveys)
###############################################################################
regional_questions = [
    'Please select your SOAF Program: ',
    'Please select your SOAP Program: ',
    'Please select your SOEA Program:',
    'Please select your SOEE Program:',
    'Please select your SOLA Program:',
    'Please select your MENA Program:',
    'Please select your SONA Program:'
]

def process_response_column(responses_df, column):
    """
    Cleans and processes a multiple-choice or rank order response column by checking types 
    and handling lists, NaNs, and other cases as in the original structure.
    
    Parameters:
    responses_df (pd.DataFrame): Survey responses DataFrame.
    column (str): Column name for the multiple-choice or rank order question.
    
    Returns:
    list: Processed list of responses with 'NULL' for NaNs.
    """
    new_list = []
    
    # Check if the column is of object type
    if responses_df[column].dtype == 'object':
        for response in responses_df[column]:
            if isinstance(response, list):
                # Add list responses directly
                new_list += response
            elif isinstance(response, float) and np.isnan(response):
                new_list.append('NULL')
            elif isinstance(response, str) and response.isdigit():
                new_list.append(response)  # Keep numeric string values
            else:
                # Log unexpected values for detailed debugging
                print(f"Unexpected value in {column}: {response} (Type: {type(response)})")
                new_list.append('NULL')
    
    elif responses_df[column].dtype in ['float', 'int64', 'int']:  # Catch int64 as well
        for response in responses_df[column]:
            if isinstance(response, float) and np.isnan(response):
                new_list.append('NULL')
            else:
                new_list.append(str(int(response)))  # Convert to int and then to string
    else:
        print(f"Problems with conversion in column {column} (Unexpected data type: {responses_df[column].dtype})")
    
    return new_list


def process_multiple_choice_question(responses_df, col, question_values_df):
    """
    Processes a multiple-choice question to generate a frequency table
    with columns ['Code', 'Value', 'Count', 'Frequency'].
    """
    new_list = process_response_column(responses_df, col)
    freq_dist = Counter(new_list)
    freq_df = pd.DataFrame.from_dict(freq_dist, orient='index').reset_index()
    freq_df.columns = ['answer_id', 'N']
    
    current_values = question_values_df[question_values_df['question_id'] == col][['answer_id', 'question_value']]
    temp_df = pd.merge(current_values, freq_df, on='answer_id', how='outer').fillna({'N': 0, 'question_value': ''})
    temp_df['N'] = temp_df['N'].astype(int)
    temp_df['Pct'] = (temp_df['N'] / len(responses_df) * 100).round(1).astype(str) + '%'
    temp_df.columns = ['Code', 'Value', 'Count', 'Frequency']
    
    return temp_df


def process_numeric_question(responses_df, col, question_values_df):
    """
    Processes a numeric question to generate a frequency table
    with columns: ['Value', 'Count', 'Frequency'].
    """
    new_list = process_response_column(responses_df, col)
    freq_dist = Counter(new_list)
    freq_df = pd.DataFrame.from_dict(freq_dist, orient='index').reset_index()
    freq_df.columns = ['answer_id', 'N']
    
    current_values = question_values_df[question_values_df['question_id'] == col][['answer_id', 'question_value']]
    temp_df = pd.merge(current_values, freq_df, on='answer_id', how='outer').fillna({'N': 0, 'question_value': ''})
    temp_df['N'] = temp_df['N'].astype(int)
    temp_df['Pct'] = (temp_df['N'] / len(responses_df) * 100).round(1).astype(str) + '%'
    
    # We rename to 'Value' (the numeric value/label), 'Count', 'Frequency'.
    temp_df.columns = ['answer_id', 'Value', 'Count', 'Frequency']
    
    # Drop 'answer_id' if you no longer need it:
    temp_df = temp_df[['Value', 'Count', 'Frequency']]
    
    return temp_df


def combine_pre_post_tables_multiple_choice(pre_df, post_df, col, question_values_df):
    """
    Combine multiple-choice question frequency tables from pre_df and post_df into:
    ['Code','Value','Pre Count','Post Count','Pre Frequency','Post Frequency'].
    """
    pre_result_df = process_multiple_choice_question(pre_df, col, question_values_df)
    post_result_df = process_multiple_choice_question(post_df, col, question_values_df)

    # Merge on ['Code','Value']
    merged_df = pd.merge(
        pre_result_df,
        post_result_df,
        how='outer',
        on=['Code','Value'],
        suffixes=('_pre','_post')
    ).fillna({'Count_pre':0,'Count_post':0,'Frequency_pre':'0.0%','Frequency_post':'0.0%'})

    merged_df = merged_df[['Code','Value','Count_pre','Count_post','Frequency_pre','Frequency_post']]
    merged_df.columns = ['Code','Value','Pre Count','Post Count','Pre Frequency','Post Frequency']
    
    return merged_df


def combine_pre_post_tables_numeric(pre_df, post_df, col, question_values_df):
    """
    Combine numeric question frequency tables from pre_df and post_df into:
    ['Value','Pre Count','Post Count','Pre Frequency','Post Frequency'].
    """
    pre_result_df = process_numeric_question(pre_df, col, question_values_df)
    post_result_df = process_numeric_question(post_df, col, question_values_df)
    
    # Merge on 'Value' only
    merged_df = pd.merge(
        pre_result_df,
        post_result_df,
        how='outer',
        on='Value',
        suffixes=('_pre','_post')
    ).fillna({'Count_pre':0,'Count_post':0,'Frequency_pre':'0.0%','Frequency_post':'0.0%'})

    merged_df = merged_df[['Value','Count_pre','Count_post','Frequency_pre','Frequency_post']]
    merged_df.columns = ['Value','Pre Count','Post Count','Pre Frequency','Post Frequency']
    
    return merged_df


def plot_pre_post_frequencies_mc(merged_df, question_title="Question Title"):
    """
    For a DataFrame with columns ['Code','Value','Pre Count','Post Count','Pre Frequency','Post Frequency'],
    create an overlapping bar chart of Pre vs. Post frequencies (based on frequency %).
    """
    pre_freq = merged_df['Pre Frequency'].str.rstrip('%').astype(float)
    post_freq = merged_df['Post Frequency'].str.rstrip('%').astype(float)
    
    x = np.arange(len(merged_df))
    width = 0.4
    
    fig, ax = plt.subplots(figsize=(7, 4.5))
    ax.bar(x - width/2, pre_freq, width=width, color='gray', alpha=0.7, label='Pre')
    ax.bar(x + width/2, post_freq, width=width, color='blue', alpha=0.7, label='Post')
    
    ax.set_xlabel('Code')
    ax.set_ylabel('Frequency (%)')
    ax.set_title(question_title)
    ax.set_xticks(x)
    ax.set_xticklabels(merged_df['Code'])
    ax.set_ylim(0, 100)
    ax.legend()
    
    # Annotate bars
    for i, (pf, p_bar) in enumerate(zip(pre_freq, ax.patches[:len(x)])):
        ax.annotate(f'{pf}%', 
                    (p_bar.get_x() + p_bar.get_width()/2, pf + 2),
                    ha='center', fontsize=9, color='gray')
    for i, (pf, p_bar) in enumerate(zip(post_freq, ax.patches[len(x):])):
        ax.annotate(f'{pf}%', 
                    (p_bar.get_x() + p_bar.get_width()/2, pf + 2),
                    ha='center', fontsize=9, color='blue')
    
    ax.grid(axis='y', linestyle='--', alpha=0.7)
    plt.tight_layout()
    return fig


def plot_pre_post_frequencies_numeric(merged_df, question_title="Numeric Question"):
    """
    For a DataFrame with columns ['Value','Pre Count','Post Count','Pre Frequency','Post Frequency'],
    create an overlapping bar chart of Pre vs. Post frequencies.
    """
    pre_freq = merged_df['Pre Frequency'].str.rstrip('%').astype(float)
    post_freq = merged_df['Post Frequency'].str.rstrip('%').astype(float)
    
    x = np.arange(len(merged_df))
    width = 0.4
    
    fig, ax = plt.subplots(figsize=(7, 4.5))
    ax.bar(x - width/2, pre_freq, width=width, color='gray', alpha=0.7, label='Pre')
    ax.bar(x + width/2, post_freq, width=width, color='blue', alpha=0.7, label='Post')
    
    ax.set_xlabel('Value')
    ax.set_ylabel('Frequency (%)')
    ax.set_title(question_title)
    ax.set_xticks(x)
    ax.set_xticklabels(merged_df['Value'])
    ax.set_ylim(0, 100)
    ax.legend()
    
    # Annotate bars
    for i, (pf, p_bar) in enumerate(zip(pre_freq, ax.patches[:len(x)])):
        ax.annotate(f'{pf}%', 
                    (p_bar.get_x() + p_bar.get_width()/2, pf + 2),
                    ha='center', fontsize=9, color='gray')
    for i, (pf, p_bar) in enumerate(zip(post_freq, ax.patches[len(x):])):
        ax.annotate(f'{pf}%', 
                    (p_bar.get_x() + p_bar.get_width()/2, pf + 2),
                    ha='center', fontsize=9, color='blue')
    
    ax.grid(axis='y', linestyle='--', alpha=0.7)
    plt.tight_layout()
    return fig


def add_pre_post_table_to_doc(doc, merged_df):
    """
    Adds the pre/post frequency table to the Word doc.
    merged_df is either:
      - MC:   Code, Value, Pre Count, Post Count, Pre Frequency, Post Frequency
      - Num:  Value, Pre Count, Post Count, Pre Frequency, Post Frequency
    """
    # Create table (header row + data rows)
    table = doc.add_table(rows=merged_df.shape[0] + 1, cols=merged_df.shape[1])
    table.style = 'Table Grid'
    
    # Header row
    for col_idx, col_name in enumerate(merged_df.columns):
        cell = table.cell(0, col_idx)
        cell.text = col_name
        cell.paragraphs[0].alignment = 1  # center
    
    # Data rows
    for row_idx in range(merged_df.shape[0]):
        for col_idx, val in enumerate(merged_df.iloc[row_idx]):
            cell = table.cell(row_idx + 1, col_idx)
            cell.text = str(val)
            # If the column is "Value," align left; else center
            if merged_df.columns[col_idx] == 'Value':
                cell.paragraphs[0].alignment = 0  # left
            else:
                cell.paragraphs[0].alignment = 1  # center

    return doc


def generate_pre_post_frequency_report(pre_df, 
                                       post_df, 
                                       question_df, 
                                       question_values_df, 
                                       regional_questions,
                                       survey_name="My Survey"):
    """
    Generates a Word doc comparing Pre vs. Post for both multiple-choice and numeric questions.
    """
    # Identify question columns by data type
    multiple_choice_columns = list(question_df.question_id[question_df["data_type"] == 'MultipleChoice'])
    numeric_choice_columns = list(question_df.question_id[question_df["data_type"] == 'Numeric'])

    doc = Document()

    # Some overview info (optional)
    doc.add_heading(f"{survey_name} - Pre vs. Post Report", level=1)
    doc.add_paragraph(f"Pre responses: N = {len(pre_df)}", style='Normal')
    doc.add_paragraph(f"Post responses: N = {len(post_df)}", style='Normal')
    doc.add_page_break()

    # Multiple-choice
    for col in multiple_choice_columns:
        qname = question_df.loc[question_df['question_id'] == col, 'question_name'].values[0]
        qtext = question_df.loc[question_df['question_id'] == col, 'question_text'].values[0]
        mapped_text = f"{qname}: {qtext}"

        # Combine
        merged_df = combine_pre_post_tables_multiple_choice(pre_df, post_df, col, question_values_df)
        
        # (Optional) Filter out rows where both Pre Count & Post Count == 0 and question is regional
        if qtext in regional_questions:
            merged_df = merged_df[(merged_df['Pre Count'] != 0) | (merged_df['Post Count'] != 0)]
        
        # Add heading
        doc.add_heading(mapped_text, level=1)

        # Insert table
        add_pre_post_table_to_doc(doc, merged_df)

        # Insert overlapping bar chart (if not too many categories)
        if len(merged_df) < 10:
            fig = plot_pre_post_frequencies_mc(merged_df, question_title=mapped_text)
            
            image_stream = BytesIO()
            plt.savefig(image_stream, format='png')
            plt.close(fig)
            image_stream.seek(0)
            
            doc.add_picture(image_stream, width=Inches(7), height=Inches(4.5))
        
        doc.add_page_break()

    # Numeric
    for col in numeric_choice_columns:
        qname = question_df.loc[question_df['question_id'] == col, 'question_name'].values[0]
        qtext = question_df.loc[question_df['question_id'] == col, 'question_text'].values[0]
        mapped_text = f"{qname}: {qtext}"

        # Combine
        merged_df = combine_pre_post_tables_numeric(pre_df, post_df, col, question_values_df)
        
        # Optional filtering for regional
        if qtext in regional_questions:
            merged_df = merged_df[(merged_df['Pre Count'] != 0) | (merged_df['Post Count'] != 0)]
        
        # Add heading
        doc.add_heading(mapped_text, level=1)

        # Insert table
        add_pre_post_table_to_doc(doc, merged_df)

        # Insert overlapping bar chart (if not too many categories)
        if len(merged_df) < 10:
            fig = plot_pre_post_frequencies_numeric(merged_df, question_title=mapped_text)
            
            image_stream = BytesIO()
            plt.savefig(image_stream, format='png')
            plt.close(fig)
            image_stream.seek(0)
            
            doc.add_picture(image_stream, width=Inches(7), height=Inches(4.5))
        
        doc.add_page_break()
    
    return doc


doc = generate_pre_post_frequency_report(
    pre_df=pre_df,
    post_df=post_df,
    question_df=question_df,
    question_values_df=question_values_df,
    regional_questions=regional_questions,
    survey_name="Lifestyle Survey"
)

doc.save("/path/to/my_pre_post_report.docx")








































